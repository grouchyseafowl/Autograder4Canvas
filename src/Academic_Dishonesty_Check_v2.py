"""
Academic Dishonesty Check v2.0
Adaptive, Data-Driven Academic Integrity Analysis System

This tool identifies markers of potential academic dishonesty using:
- Externalized, versioned detection markers (YAML configuration)
- Peer comparison (statistical outlier detection)
- Context-aware adjustments (ESL, first-gen, community college)
- Pedagogically-informed reporting (conversation starters, not accusations)

IMPORTANT: This tool is a CONVERSATION STARTER, not a verdict.
Flags are indicators for instructor judgment, not proof of dishonesty.
"""

import os
import re
import sys
import json
import platform
import statistics
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from collections import defaultdict
from dataclasses import dataclass, field
from html import unescape

# Add current directory to path for module imports
sys.path.insert(0, str(Path(__file__).parent))

# Import context analyzer for ESL detection
try:
    from modules.context_analyzer import ContextAnalyzer, StudentContext
    HAS_CONTEXT_ANALYZER = True
except ImportError:
    HAS_CONTEXT_ANALYZER = False
    print("⚠ Warning: Context analyzer module not available. ESL detection disabled.")

# Import two-axis weight composer (Phase 8: per-marker calibration system)
try:
    from modules.weight_composer import ComposedWeights, PopulationSettings
    HAS_WEIGHT_COMPOSER = True
except ImportError:
    HAS_WEIGHT_COMPOSER = False
    ComposedWeights = None  # type: ignore
    PopulationSettings = None  # type: ignore

# Import feedback tracker for instructor feedback
try:
    from modules.feedback_tracker import FeedbackTracker, FeedbackRecord, print_feedback_summary
    HAS_FEEDBACK_TRACKER = True
except ImportError:
    HAS_FEEDBACK_TRACKER = False
    print("⚠ Warning: Feedback tracker module not available.")

# Import organizational analyzer for AI-specific pattern detection
try:
    from modules.organizational_analyzer import OrganizationalAnalyzer
    HAS_ORG_ANALYZER = True
except ImportError:
    HAS_ORG_ANALYZER = False
    print("⚠ Warning: organizational_analyzer not available - AI-specific organizational detection disabled")

# Import human presence detector for paradigm-shift detection
try:
    from modules.human_presence_detector import HumanPresenceDetector
    HAS_HUMAN_PRESENCE = True
except ImportError:
    HAS_HUMAN_PRESENCE = False
    print("⚠ Warning: human_presence_detector not available - Human presence detection disabled")

# PHASE 7: Import assignment configuration system
try:
    from modules.assignment_config import (
        AssignmentConfigLoader,
        apply_assignment_config_to_detector
    )
    HAS_ASSIGNMENT_CONFIG = True
except ImportError:
    HAS_ASSIGNMENT_CONFIG = False
    print("⚠ Warning: assignment_config not available - Using default settings")

# Import gibberish gate for pre-analysis filtering
try:
    from insights.gibberish_gate import check_gibberish, GibberishResult
    HAS_GIBBERISH_GATE = True
except ImportError:
    HAS_GIBBERISH_GATE = False

# Import citation checker for engagement signal
try:
    from insights.citation_checker import extract_citations
    HAS_CITATION_CHECKER = True
except ImportError:
    HAS_CITATION_CHECKER = False

# Import cohort calibration (Mechanism 1: class-relative baselines)
try:
    from modules.cohort_calibration import (
        CohortCalibrator, extract_signal_vector, SIGNAL_NAMES
    )
    HAS_COHORT_CALIBRATION = True
except ImportError:
    HAS_COHORT_CALIBRATION = False

# Version info
VERSION = "2.0.0"
VERSION_DATE = "2025-12-26"

# =============================================================================
# CONFIGURATION
# =============================================================================

CANVAS_BASE_URL = os.getenv("CANVAS_BASE_URL")
API_TOKEN = os.getenv("CANVAS_API_TOKEN")

# Try to import requests
try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

HEADERS = {}
if API_TOKEN and HAS_REQUESTS:
    HEADERS = {
        "Authorization": f"Bearer {API_TOKEN}",
        "Content-Type": "application/json"
    }


def get_config_dir() -> Path:
    """Get the configuration directory for storing settings."""
    system = platform.system()
    
    if system == "Windows":
        base = Path(os.environ.get("LOCALAPPDATA", Path.home() / "AppData" / "Local"))
        config_dir = base / "CanvasAutograder"
    elif system == "Darwin":
        config_dir = Path.home() / "Library" / "Application Support" / "CanvasAutograder"
    else:
        config_dir = Path.home() / ".config" / "CanvasAutograder"
    
    # Security: Restrict permissions to owner only (user read/write/execute)
    config_dir.mkdir(parents=True, exist_ok=True, mode=0o700)
    return config_dir


def get_output_base_dir() -> Path:
    """Get base output directory for reports."""
    system = platform.system()
    
    # Check for /output directory (container environment)
    if os.path.isdir("/output"):
        return Path("/output")
    
    # Check for custom output directory in config
    config_file = get_config_dir() / "settings.json"
    if config_file.exists():
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
                if "output_directory" in config:
                    custom_dir = Path(config["output_directory"])
                    custom_dir.mkdir(parents=True, exist_ok=True)
                    return custom_dir
        except Exception:
            pass
    
    # Default location
    if system == "Windows":
        documents = Path(os.environ.get("USERPROFILE", Path.home())) / "Documents"
    else:
        documents = Path.home() / "Documents"
    
    return documents / "Autograder Rationales"


# =============================================================================
# UNICODE NORMALIZATION
# =============================================================================

# Zero-width and invisible characters that could be used for text manipulation
_ZERO_WIDTH_CHARS = re.compile(
    '[\u200b\u200c\u200d\u200e\u200f'   # zero-width space/joiner/non-joiner/marks
    '\u2060\u2061\u2062\u2063\u2064'     # word joiner, invisible operators
    '\ufeff'                              # byte order mark
    '\u00ad'                              # soft hyphen
    '\u034f'                              # combining grapheme joiner
    '\u061c'                              # Arabic letter mark
    '\u115f\u1160'                        # Hangul fillers
    '\u17b4\u17b5'                        # Khmer invisible
    '\u180e'                              # Mongolian vowel separator
    '\uffa0'                              # halfwidth Hangul filler
    ']'
)


def _detect_unicode_manipulation(text: str) -> Tuple[int, List[str]]:
    """Detect zero-width / invisible characters BEFORE stripping them.

    This is a Tier 1 integrity signal — invisible-character insertion is
    active gaming behavior (circumvention, not ambiguity).

    Returns:
        (count_found, detail_strings) — count of invisible chars and
        human-readable descriptions of what was found.
    """
    matches = _ZERO_WIDTH_CHARS.findall(text)
    if not matches:
        return 0, []

    import unicodedata
    char_names = {}
    for ch in matches:
        name = unicodedata.name(ch, f'U+{ord(ch):04X}')
        char_names[name] = char_names.get(name, 0) + 1

    details = [f"{count}x {name}" for name, count in char_names.items()]
    return len(matches), details


def _normalize_unicode(text: str) -> str:
    """Strip zero-width characters and normalize whitespace.

    Catches invisible-character insertion (a gaming technique where students
    embed zero-width characters to break pattern detection or inflate word
    count).  Detection is handled separately by _detect_unicode_manipulation();
    this function normalizes so the actual text can be analyzed accurately.
    """
    import unicodedata
    # NFC normalization first (canonical composition)
    text = unicodedata.normalize('NFC', text)
    # Strip zero-width / invisible characters
    text = _ZERO_WIDTH_CHARS.sub('', text)
    # Normalize whitespace (multiple spaces → single, strip leading/trailing)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


# =============================================================================
# BUILT-IN PATTERNS
# =============================================================================

# AI detection patterns - transitions and meta-commentary
AI_TRANSITIONS = [
    "it is important to note", "it should be noted", "it is worth noting",
    "delving into", "navigating the complexities", "in the realm of",
    "this underscores", "it becomes evident", "one cannot overstate",
    "serves as a testament", "in light of", "furthermore", "moreover",
    "additionally", "in addition", "in conclusion", "to sum up",
    "firstly", "secondly", "thirdly", "lastly", "this essay will explore",
    "this paper will examine", "as previously mentioned", "as stated above",
    "it is crucial to understand", "plays a pivotal role", "a profound impact"
]

# Generic/vague phrases
GENERIC_PHRASES = [
    "throughout history", "since the beginning of time", "in today's society",
    "in today's world", "many things", "various aspects", "studies show",
    "research indicates", "experts say", "it has been shown", "some people believe",
    "a variety of", "a number of", "plays a crucial role", "of paramount importance",
    "cannot be overstated", "it can be said", "it could be argued",
    "in many ways", "on the other hand", "when it comes to"
]

# Personal voice markers (absence is suspicious in personal writing)
PERSONAL_MARKERS = [
    r"\bI\s", r"\bmy\b", r"\bme\b", r"\bmine\b", r"\bmyself\b",
    r"\bwe\b", r"\bour\b", "I realized", "I remember", "I noticed",
    "my experience", "I believe", "I think", "I felt", "I was"
]

# Emotional language
EMOTIONAL_MARKERS = [
    "I felt", "I was scared", "I cried", "I laughed", "it hurt",
    "I struggled", "I was confused", "I was frustrated", "I was excited",
    "I was nervous", "I was embarrassed", "I was proud", "I was relieved"
]

# Cognitive diversity markers (protective indicators of neurodivergent engagement)
COGNITIVE_DIVERSITY_MARKERS = [
    # Depth over breadth / Hyperfocus
    "what really interests me", "I spent a lot of time thinking about",
    "I went down a rabbit hole", "I could talk about this for hours",
    "this is something I've been interested in", "I know more than I probably need to",
    # Associative thinking
    "this reminds me of something completely different", "I know this is a tangent",
    "this might seem unrelated but", "going back to", "now that I think about it",
    "wait, this is like", "I just realized", "I see a pattern",
    # Authentic struggle
    "I'm not sure how to organize this", "I'm having trouble explaining",
    "I keep losing my train of thought", "let me try that again",
    "I know I'm jumping around", "this is harder to explain than I thought",
    "trying to make this make sense", "let me start over",
    # Precise/literal engagement
    "to be specific", "to be precise", "more accurately", "technically",
    "I mean exactly", "I'm being literal", "the exact wording",
    # Metacognitive commentary
    "I'm realizing as I write", "thinking about it now", "my brain keeps going to",
    "the way I understand this", "my thought process", "as I'm thinking through this"
]

# Inflated vocabulary
INFLATED_VOCAB = [
    ("utilize", "use"), ("demonstrate", "show"), ("individuals", "people"),
    ("commence", "start"), ("terminate", "end"), ("endeavor", "try"),
    ("facilitate", "help"), ("implement", "do"), ("ascertain", "find out"),
    ("optimal", "best"), ("subsequent", "next"), ("prior to", "before"),
    ("multifaceted", "complex"), ("plethora", "many")
]


# =============================================================================
# ASSIGNMENT PROFILES
# =============================================================================

ASSIGNMENT_PROFILES = {
    "personal_reflection": {
        "id": "personal_reflection",
        "name": "Personal Reflection / Response Paper",
        "description": "Personal essays, reflections requiring authentic voice",
        "detection_focus": "absence_detection",
        "weight_multipliers": {
            "personal_voice": 2.0,
            "specific_details": 2.0,
            "emotional_language": 1.5,
            "ai_transitions": 1.3,
            "generic_phrases": 1.5
        },
        "instructor_notes": [
            "Absence of personal voice is as concerning as presence of AI markers",
            "Look for specific sensory details and named individuals/places",
            "Generic reflections that could apply to anyone are concerning"
        ]
    },
    "analytical_essay": {
        "id": "analytical_essay",
        "name": "Analytical / Argumentative Essay",
        "description": "Formal analysis essays where some formal language is expected",
        "detection_focus": "presence_detection",
        "weight_multipliers": {
            "personal_voice": 0.5,
            "ai_transitions": 0.8,
            "generic_phrases": 1.2,
            "citation_markers": 1.5
        },
        "instructor_notes": [
            "Some formal transitions are expected in analytical writing",
            "Focus on generic content that could be written without reading the texts",
            "Check if analysis engages with specific passages"
        ]
    },
    "discussion_post": {
        "id": "discussion_post",
        "name": "Discussion Forum Post",
        "description": "Forum posts that should reference readings and classmates",
        "detection_focus": "engagement_detection",
        "weight_multipliers": {
            "personal_voice": 1.0,
            "ai_transitions": 1.2,
            "generic_phrases": 1.3,
            "engagement_markers": 1.5
        },
        "instructor_notes": [
            "Posts should reference specific readings or classmate comments",
            "Highly formal language is unusual in discussion posts",
            "Generic posts that don't engage with course material are concerning"
        ]
    },
    "rough_draft": {
        "id": "rough_draft",
        "name": "Rough Draft / First Draft",
        "description": "Early drafts where errors and rough structure are expected",
        "detection_focus": "polish_detection",
        "weight_multipliers": {
            "personal_voice": 0.5,
            "ai_transitions": 0.7,
            "generic_phrases": 0.8,
            "polish_markers": 2.0
        },
        "instructor_notes": [
            "Drafts should show rough edges - too-perfect drafts are suspicious",
            "Students often ignore 'don't polish' instructions, so focus on patterns",
            "Compare against final version for authentic revision evidence"
        ]
    },
    "standard": {
        "id": "standard",
        "name": "Standard Analysis (Default)",
        "description": "Default balanced analysis for any assignment type",
        "detection_focus": "balanced",
        "weight_multipliers": {
            "personal_voice": 1.0,
            "ai_transitions": 1.0,
            "generic_phrases": 1.0
        },
        "instructor_notes": [
            "Default mode - use more specific profiles when possible",
            "All checks enabled with moderate sensitivity"
        ]
    }
}


# =============================================================================
# DATA CLASSES
# =============================================================================

@dataclass
class AnalysisResult:
    """Result of analyzing a single submission."""
    student_id: str
    student_name: str
    
    # Raw text info
    text: str
    word_count: int
    
    # Marker scores
    suspicious_score: float
    authenticity_score: float
    marker_counts: Dict[str, int]
    markers_found: Dict[str, List[str]]
    
    # Concern assessment
    concern_level: str
    
    # Context adjustments
    context_adjustments_applied: List[str] = field(default_factory=list)
    adjusted_suspicious_score: Optional[float] = None
    adjusted_concern_level: Optional[str] = None
    
    # Peer comparison (filled in later)
    suspicious_percentile: Optional[float] = None
    authenticity_percentile: Optional[float] = None
    is_outlier: bool = False
    outlier_reasons: List[str] = field(default_factory=list)
    
    # AI-specific organizational analysis
    organizational_analysis: Optional[Dict] = None
    ai_organizational_score: float = 0.0

    # Human presence detection (paradigm shift)
    human_presence_confidence: Optional[float] = None
    human_presence_level: Optional[str] = None
    human_presence_details: Optional[Dict] = None

    # Student-provided context (PHASE 6: Privacy-focused, optional)
    student_context: Optional[str] = None
    student_context_applied: bool = False
    student_context_adjustments: List[str] = field(default_factory=list)

    # Smoking gun: raw chatbot artifacts (HTML/markdown paste)
    # Detected BEFORE text cleaning — evidence is destroyed by HTML stripping.
    # When True, concern_level is forced to 'high' regardless of other scores.
    smoking_gun: bool = False
    smoking_gun_details: List[str] = field(default_factory=list)

    # Tier 1 integrity signals (engagement reframe)
    # These are binary, non-negotiable flags — no cultural bias risk.
    unicode_manipulation: bool = False
    unicode_manipulation_details: List[str] = field(default_factory=list)
    gibberish_detected: bool = False
    gibberish_reason: str = ""
    gibberish_detail: str = ""

    # Citation analysis (engagement signal)
    citation_count: int = 0
    specific_citation_count: int = 0
    generic_reading_ref_count: int = 0
    citations: List[Any] = field(default_factory=list)

    # Engagement reframe signals (Phase A)
    # These measure HOW the student engaged, not WHETHER they cheated.
    engagement_signals: Optional[Dict[str, Any]] = None

    # Guidance
    conversation_starters: List[str] = field(default_factory=list)
    revision_guidance: List[str] = field(default_factory=list)
    verification_questions: List[str] = field(default_factory=list)

    # Phase 8: Two-axis weight system provenance
    education_level: Optional[str] = None       # which edu level profile was used
    population_settings: Optional[Dict] = None  # effective population settings dict

    # Mechanism 1: Cohort calibration — class-relative signal interpretations
    # Maps signal_name → interpretation string (e.g., 'typical', 'conversation_opportunity')
    cohort_percentiles: Optional[Dict[str, str]] = None


# =============================================================================
# ANALYSIS ENGINE
# =============================================================================

class DishonestyAnalyzer:
    """
    Main analysis engine for academic dishonesty detection.
    
    Combines:
    - Marker-based detection (built-in patterns)
    - Peer comparison (statistical outlier detection)
    - Context-aware adjustments
    - Pedagogical reporting
    """
    
    def __init__(self,
                 profile_id: str = "standard",
                 context_profile: str = "community_college",
                 assignment_type: Optional[str] = None,
                 course_level: Optional[str] = None,
                 institutional_context: Optional[str] = None,
                 composed_weights=None,
                 aic_config: Optional[dict] = None):
        """
        Initialize the analyzer.

        Args:
            profile_id: Assignment profile to use (legacy)
            context_profile: Student population context (legacy)
            assignment_type: PHASE 7 - Assignment type (legacy yaml path)
            course_level: PHASE 7 - Course level
            institutional_context: PHASE 7 - Institution type
            composed_weights: PHASE 8 - ComposedWeights from WeightComposer.
            aic_config: PHASE 9 - Template-sourced AIC config dict (from
                        assignment_templates.get_aic_config). When provided,
                        overrides Phase 7 yaml flags and Phase 8 marker weights
                        with the teacher-configured values stored in the template.
                        Keys: aic_mode, personal_voice_authentic, invert_sentence_signals,
                        weight_personal_voice, weight_ai_patterns, weight_course_content,
                        weight_rough_work.
        """
        self.profile_id = profile_id
        self.context_profile = context_profile
        self.profile = ASSIGNMENT_PROFILES.get(profile_id, ASSIGNMENT_PROFILES["standard"])

        # PHASE 7: Assignment configuration
        self.assignment_type = assignment_type
        self.course_level = course_level
        self.institutional_context = institutional_context or context_profile  # Backward compatible

        # Load assignment configuration
        # v3.0: also load mode flags (personal_voice_authentic, invert_sentence_signals)
        self._personal_voice_authentic = True   # default: personal voice = authenticity signal
        self._invert_sentence_signals = False   # default: smooth prose = suspicious (normal polarity)
        if HAS_ASSIGNMENT_CONFIG:
            self.config_loader = AssignmentConfigLoader()
            # Get multipliers from course level and institutional context
            self.config_multiplier, self.authenticity_boost = \
                self.config_loader.get_combined_multiplier(course_level, self.institutional_context)
            # Load mode-specific signal flags
            if assignment_type and assignment_type != 'auto':
                pva, iss = self.config_loader.get_mode_flags(assignment_type)
                self._personal_voice_authentic = pva
                self._invert_sentence_signals = iss
        else:
            self.config_loader = None
            self.config_multiplier = 1.0
            self.authenticity_boost = 1.0

        # Context multiplier (community college = more lenient) - LEGACY
        legacy_multiplier = 0.7 if context_profile == "community_college" else 1.0

        # Combine legacy and new multipliers
        self.context_multiplier = legacy_multiplier * self.config_multiplier

        # PHASE 8: Two-axis weight system
        self.composed_weights = composed_weights
        if composed_weights is not None:
            # Per-marker weights from WeightComposer replace hardcoded defaults
            self._marker_weights = composed_weights.marker_weights
            self._cognitive_protection_floor = composed_weights.cognitive_protection_floor
        else:
            # Hardcoded defaults (backward-compat path — identical to pre-Phase 8 behavior)
            self._marker_weights = {
                'ai_transitions': 0.5,
                'generic_phrases': 0.4,
                'inflated_vocabulary': 0.3,
                'ai_specific_org': 1.0,
                'personal_voice': 0.5,
                'emotional_language': 0.8,
                'cognitive_diversity': 0.6,
                'hp_absence': 0.5,
            }
            self._cognitive_protection_floor = 0.5

        # PHASE 9: Template-sourced AIC config (highest priority — overrides all above)
        # aic_config comes from assignment_templates.get_aic_config(template).
        # It carries the teacher's saved weight values; ignore yaml path entirely.
        self._aic_config = aic_config
        if aic_config is not None:
            # Signal polarity flags — override Phase 7 yaml values
            self._personal_voice_authentic = aic_config.get(
                "personal_voice_authentic", self._personal_voice_authentic
            )
            self._invert_sentence_signals = aic_config.get(
                "invert_sentence_signals", self._invert_sentence_signals
            )
            # Apply weight multipliers to _marker_weights
            # (skip if composed_weights is active — Phase 8 already owns those)
            if composed_weights is None:
                w_pv = aic_config.get("weight_personal_voice", 1.0)
                w_ai = aic_config.get("weight_ai_patterns",  1.0)
                if w_pv != 1.0:
                    self._marker_weights['personal_voice'] = (
                        self._marker_weights.get('personal_voice', 0.5) * w_pv
                    )
                    self._marker_weights['emotional_language'] = (
                        self._marker_weights.get('emotional_language', 0.8) * w_pv
                    )
                if w_ai != 1.0:
                    self._marker_weights['ai_transitions'] = (
                        self._marker_weights.get('ai_transitions', 0.5) * w_ai
                    )
                    self._marker_weights['generic_phrases'] = (
                        self._marker_weights.get('generic_phrases', 0.4) * w_ai
                    )

    def analyze_text(self,
                     text: str,
                     student_id: str = "unknown",
                     student_name: str = "Unknown Student",
                     student_context: Optional[str] = None) -> AnalysisResult:
        """
        Analyze a single text submission.

        Args:
            text: The student's text to analyze
            student_id: Student identifier
            student_name: Student name
            student_context: Optional context from student about their writing process
                           (e.g., "I tend to write in a non-linear way" or "English is my second language")
        """
        # Detect raw AI artifacts BEFORE cleaning destroys HTML/markdown evidence
        smoking_gun, smoking_gun_details, artifact_counts = \
            self._detect_raw_ai_artifacts(text)

        # Tier 1 integrity signal: detect unicode manipulation BEFORE stripping
        unicode_manipulation_count, unicode_manipulation_details = \
            _detect_unicode_manipulation(text)

        # Strip zero-width characters and normalize unicode before analysis
        text = _normalize_unicode(text)

        # Clean text
        text = self._clean_text(text)
        word_count = len(text.split())

        # Tier 1 integrity signal: gibberish gate (incoherence check)
        # Catches keyboard mash, Lorem Ipsum, repetition spam BEFORE
        # expensive analysis runs. Conservative — only flags obvious cases.
        gibberish_result = None
        if HAS_GIBBERISH_GATE:
            try:
                gibberish_result = check_gibberish(text)
            except Exception:
                pass

        # Engagement signal: citation extraction
        # Distinguishes specific citations (author-year, URL, DOI, quoted title)
        # from generic references ("the reading says...").
        # Source depth is a Tier 3 engagement signal per spec.
        citations = []
        specific_citation_count = 0
        generic_reading_ref_count = 0
        if HAS_CITATION_CHECKER:
            try:
                citations = extract_citations(
                    text, student_id=student_id, student_name=student_name
                )
                specific_citation_count = sum(
                    1 for c in citations
                    if c.citation_type != 'reading_reference'
                )
                generic_reading_ref_count = sum(
                    1 for c in citations
                    if c.citation_type == 'reading_reference'
                )
            except Exception:
                pass

        # Analyze with built-in patterns
        suspicious_score, authenticity_score, marker_counts, markers_found, \
            cognitive_protection_multiplier = self._analyze_with_builtin_patterns(text)
        # Save pattern-only suspicious for convergence channel check (Phase 2)
        pattern_suspicious = suspicious_score

        # Apply profile weight multipliers (legacy path only)
        # When ComposedWeights are active, per-marker weights are already correct —
        # skip _apply_profile_weights which would recalculate from base_weight=0.3
        if self.composed_weights is None:
            suspicious_score = self._apply_profile_weights(suspicious_score, marker_counts)

        # Analyze organizational patterns (AI-specific signatures)
        # These are patterns that don't overlap with neurodivergent writing
        org_analysis = None
        ai_org_score = 0.0

        if HAS_ORG_ANALYZER:
            try:
                org_analyzer = OrganizationalAnalyzer()
                org_analysis = org_analyzer.analyze(text)
                ai_org_score = org_analysis.total_ai_organizational_score

                # Add to suspicious score (scaled by ai_specific_org weight)
                # NOTE: This score is NOT subject to cognitive diversity protection
                # (except via the ai_specific_org weight in ComposedWeights ND-aware mode)
                ai_org_weight = self._marker_weights.get('ai_specific_org', 1.0)
                suspicious_score += ai_org_score * ai_org_weight

                # Track in marker counts for transparency
                if ai_org_score > 0:
                    marker_counts['ai_specific_organization'] = 1  # Binary indicator

                    # Build detailed description for transparency
                    details = []
                    if org_analysis.excessive_headers:
                        details.append(f"excessive headers ({org_analysis.details['header_analysis']['total_count']})")
                    if org_analysis.hierarchical_headers:
                        details.append(f"deep hierarchy ({org_analysis.details['header_analysis']['deepest_level']} levels)")
                    if org_analysis.balanced_sections:
                        variance = org_analysis.details['section_analysis']['variance_coefficient']
                        details.append(f"balanced sections (variance={variance:.2f})")
                    if org_analysis.uniform_paragraphs:
                        variance = org_analysis.details['paragraph_analysis']['variance_coefficient']
                        details.append(f"uniform paragraphs (variance={variance:.2f})")

                    markers_found['ai_specific_organization'] = [
                        f"AI organizational signature (score={ai_org_score:.1f}): {', '.join(details)}"
                    ]

            except Exception as e:
                print(f"  ⚠ Organizational analysis skipped: {e}")
                import traceback
                traceback.print_exc()

        # Detect human presence (paradigm shift: presence-based detection)
        human_presence_result = None
        human_confidence = None
        human_level = None

        if HAS_HUMAN_PRESENCE:
            try:
                hp_detector = HumanPresenceDetector()

                # PHASE 7: Apply assignment configuration to detector
                if HAS_ASSIGNMENT_CONFIG and self.assignment_type:
                    apply_assignment_config_to_detector(
                        hp_detector,
                        assignment_type=self.assignment_type,
                        course_level=self.course_level,
                        institutional_context=self.institutional_context
                    )

                # PHASE 9: Apply template weight overrides to HPD category weights
                # weight_course_content → contextual_grounding
                # weight_rough_work    → cognitive_struggle + productive_messiness
                if self._aic_config is not None:
                    _w_cc = self._aic_config.get("weight_course_content", 1.0)
                    _w_rw = self._aic_config.get("weight_rough_work", 1.0)
                    if (_w_cc != 1.0 or _w_rw != 1.0) and hasattr(hp_detector, 'CATEGORY_WEIGHTS'):
                        if not hasattr(hp_detector, '_original_weights'):
                            hp_detector._original_weights = hp_detector.CATEGORY_WEIGHTS.copy()
                        if _w_cc != 1.0 and 'contextual_grounding' in hp_detector.CATEGORY_WEIGHTS:
                            hp_detector.CATEGORY_WEIGHTS['contextual_grounding'] = (
                                hp_detector._original_weights['contextual_grounding'] * _w_cc
                            )
                        if _w_rw != 1.0:
                            for _k in ('cognitive_struggle', 'productive_messiness'):
                                if _k in hp_detector.CATEGORY_WEIGHTS:
                                    hp_detector.CATEGORY_WEIGHTS[_k] = (
                                        hp_detector._original_weights[_k] * _w_rw
                                    )
                        _total = sum(hp_detector.CATEGORY_WEIGHTS.values())
                        if _total > 0:
                            for _k in hp_detector.CATEGORY_WEIGHTS:
                                hp_detector.CATEGORY_WEIGHTS[_k] /= _total

                # Analyze with configured weights
                human_presence_result = hp_detector.analyze(
                    text,
                    assignment_type=self.assignment_type or self.profile_id
                )
                human_confidence = human_presence_result.confidence_percentage
                human_level = human_presence_result.confidence_level
            except Exception as e:
                print(f"  ⚠ Human presence analysis skipped: {e}")
                import traceback
                traceback.print_exc()

        # Detect ESL patterns (strong indicator of human authorship)
        esl_detected = False
        esl_adjustment = 1.0
        esl_notes = []  # accumulated here, merged into context_applied later
        if HAS_CONTEXT_ANALYZER:
            try:
                context_analyzer = ContextAnalyzer()
                context_result = context_analyzer.analyze_context(text)

                # If ESL error patterns detected, reduce suspicious score
                if context_result.context.has_esl_error_patterns:
                    esl_detected = True
                    # ESL errors are strong evidence of HUMAN authorship
                    # AI models don't make these mistakes
                    esl_adjustment = 0.6  # 40% reduction in suspicious score
                    authenticity_score += 2.0  # Boost authenticity

                    # Zero bias-carrying signals for ESL students:
                    # Comma density: formal ESL writing uses subordinate clauses
                    #   learned in grammar instruction — a linguistic ASSET, not
                    #   a disengagement signal (#COMMUNITY_CULTURAL_WEALTH)
                    # Avg word length: correlates with education resources, not
                    #   engagement (#ETHNIC_STUDIES — word gap = resource gap)
                    # Starter diversity: careful L2 construction may produce high
                    #   diversity without AI (#LANGUAGE_JUSTICE)
                    if org_analysis:
                        sent = org_analysis.details.get('sentence_analysis', {})
                        esl_zeroed = (
                            sent.get('comma_density_score', 0) +
                            sent.get('avg_word_length_score', 0) +
                            sent.get('starter_diversity_score', 0)
                        )
                        if esl_zeroed > 0:
                            ai_org_weight = self._marker_weights.get('ai_specific_org', 1.0)
                            suspicious_score -= esl_zeroed * ai_org_weight
                            suspicious_score = max(suspicious_score, 0.0)
                            esl_notes.append(
                                f"ESL: zeroed structural signals "
                                f"(-{esl_zeroed * ai_org_weight:.2f})"
                            )
            except Exception as e:
                # Don't let ESL detection failure break analysis
                print(f"  ⚠ ESL detection skipped: {e}")

        # Linguistic feature detection (subsumes ESL, adds AAVE/affect/convention)
        # Runs alongside ContextAnalyzer — the new detector covers everything the
        # old one did plus additional features.  AIC weight adjustments are applied
        # multiplicatively to self._marker_weights so both systems stack.
        try:
            from modules.linguistic_features import detect_features
            _repertoire = detect_features(text, word_count, was_translated=False)
            if _repertoire.aic_adjustments:
                for _lf_marker, _lf_mult in _repertoire.aic_adjustments.items():
                    if _lf_marker in self._marker_weights:
                        self._marker_weights[_lf_marker] *= _lf_mult
                        esl_notes.append(
                            f"linguistic_features: {_lf_marker} *= {_lf_mult:.2f}"
                        )
        except ImportError:
            pass  # graceful fallback — module not yet available
        except Exception as e:
            print(f"  ⚠ Linguistic feature detection skipped: {e}")

        # PHASE 6: Process optional student-provided context
        student_context_applied = False
        student_context_adjustments = []
        if student_context:
            context_adj_susp, context_adj_org, adjustments = self._process_student_context(
                student_context, suspicious_score, ai_org_score
            )

            if adjustments:
                student_context_applied = True
                student_context_adjustments = adjustments

                # Apply student context adjustments
                suspicious_score = context_adj_susp
                ai_org_score = context_adj_org

        # PHASE 7: Apply authenticity boost from assignment configuration
        if HAS_ASSIGNMENT_CONFIG and self.authenticity_boost != 1.0:
            authenticity_score *= self.authenticity_boost

        # ── Human presence bridge (OBS-AIC-01) ───────────────────────────
        # Low human presence in modes expecting personal voice → suspicious.
        # Mode-sensitive via weight_personal_voice from template system.
        if human_confidence is not None:
            hp_absence_threshold = 20.0
            hp_absence_weight = self._marker_weights.get('hp_absence', 0.5)

            # Mode sensitivity: discussion/personal modes weight this more
            if self._aic_config is not None:
                hp_mode_mult = self._aic_config.get('weight_personal_voice', 1.0)
            else:
                hp_mode_mult = 1.0 if self._personal_voice_authentic else 0.3

            if human_confidence < hp_absence_threshold:
                absence_signal = (hp_absence_threshold - human_confidence) / hp_absence_threshold
                absence_contribution = absence_signal * hp_absence_weight * hp_mode_mult * 1.0
                absence_contribution = min(absence_contribution, 0.75)
                # Cognitive diversity protection
                absence_contribution *= cognitive_protection_multiplier
                suspicious_score += absence_contribution

                if absence_contribution > 0.1:
                    marker_counts['hp_absence'] = 1
                    markers_found['hp_absence'] = [
                        f"Low human presence ({human_confidence:.1f}%) — "
                        f"contribution: {absence_contribution:.2f} "
                        f"(mode×{hp_mode_mult:.1f}, cog_prot×{cognitive_protection_multiplier:.1f})"
                    ]

        # Apply context adjustments (ESL, community college, etc.)
        # When ComposedWeights are active: adjustments baked into per-marker weights;
        # context_multiplier is skipped. ESL error pattern adjustment also skipped
        # (population-level ESL overlay in ComposedWeights handles calibration).
        if self.composed_weights is not None:
            adjusted_suspicious = suspicious_score
        else:
            adjusted_suspicious = suspicious_score * self.context_multiplier * esl_adjustment
        
        # ── Signal convergence (OBS-AIC-08) ─────────────────────────────
        # When 3+ independent channels agree, amplify suspicious score.
        # Channels checked against pattern-only suspicious (before org/HP added).
        converging = 0
        convergence_channels = []

        if pattern_suspicious > 0.3:
            converging += 1
            convergence_channels.append('pattern_markers')
        if ai_org_score > 0.0:
            converging += 1
            convergence_channels.append('organizational')
        if human_confidence is not None and human_confidence < 15.0:
            converging += 1
            convergence_channels.append('hp_absence')
        if authenticity_score < 2.0:
            converging += 1
            convergence_channels.append('authenticity_deficit')
        # Sentence uniformity from org analysis (gradient zone: < 0.35)
        if org_analysis:
            sent_analysis = org_analysis.details.get('sentence_analysis', {})
            sent_vc = sent_analysis.get('variance_coefficient', 1.0)
            if sent_vc < 0.35:
                converging += 1
                convergence_channels.append('uniformity')
            # Starter diversity: use SCORED value (requires 8+ sentences
            # to avoid small-sample false positives — short ESL texts with
            # 5 formal sentences achieve 1.0 by chance)
            # (#LANGUAGE_JUSTICE, #ALGORITHMIC_JUSTICE)
            starter_div_score = sent_analysis.get('starter_diversity_score', 0.0)
            if starter_div_score > 0:
                converging += 1
                convergence_channels.append('starter_diversity')
            # Comma density + avg word length: CORROBORATION ONLY.
            # Skip entirely when ESL detected — these signals carry
            # population-dependent bias risk (#LANGUAGE_JUSTICE).
            if not esl_detected:
                comma_d_score = sent_analysis.get('comma_density_score', 0.0)
                if comma_d_score > 0:
                    converging += 1
                    convergence_channels.append('comma_density')
                awl_score = sent_analysis.get('avg_word_length_score', 0.0)
                if awl_score > 0:
                    converging += 1
                    convergence_channels.append('avg_word_length')

        if converging >= 3 and suspicious_score > 0:
            convergence_multiplier = 1.0 + 0.2 * (converging - 2)
            suspicious_score *= convergence_multiplier
            marker_counts['signal_convergence'] = converging
            markers_found['signal_convergence'] = [
                f"{converging} channels: {', '.join(convergence_channels)} "
                f"— {convergence_multiplier:.1f}x"
            ]

        # Determine concern level
        concern_level = self._determine_concern_level(
            suspicious_score, authenticity_score, word_count=word_count,
        )
        adjusted_concern = self._determine_concern_level(
            adjusted_suspicious, authenticity_score, word_count=word_count,
        )
        
        # Get guidance
        conversation_starters = self._get_conversation_starters(concern_level)
        revision_guidance = self._get_revision_guidance(concern_level)
        verification_questions = self._get_verification_questions(concern_level)
        
        context_applied = []
        if self.composed_weights is not None:
            # Phase 8: Two-axis weight system active — log composition summary
            cw = self.composed_weights
            context_applied.append(
                f"Two-axis weight system active: "
                f"education_level={cw.education_level}, "
                f"esl={cw.population.esl_level}, "
                f"first_gen={cw.population.first_gen_level}, "
                f"nd_aware={cw.population.neurodivergent_aware}"
            )
        elif self.context_profile == "community_college":
            context_applied = ["Community college population adjustment (30% more lenient)"]
        if esl_detected:
            context_applied.append("ESL error patterns detected - strong indicator of human authorship")
            context_applied.append("Note: AI models don't make article errors or tense mixing")
            context_applied.extend(esl_notes)
        if student_context_applied:
            context_applied.extend(student_context_adjustments)

        # Brevity cap note — if word_count was below the reliable-analysis
        # threshold, _determine_concern_level already capped at 'low'.
        # Surface that decision in context_adjustments so it's transparent.
        if word_count < self.MIN_WORDS_FOR_RELIABLE_ANALYSIS:
            context_applied.append(
                f"Short submission ({word_count} words, threshold "
                f"{self.MIN_WORDS_FOR_RELIABLE_ANALYSIS}): text too brief for "
                f"reliable authenticity analysis — concern level capped at 'low'"
            )

        # Add organizational analysis details if present
        if ai_org_score > 0:
            org_details = []
            if org_analysis.excessive_headers:
                count = org_analysis.details['header_analysis']['total_count']
                org_details.append(f"Excessive headers detected ({count} headers)")

            if org_analysis.balanced_sections:
                variance = org_analysis.details['section_analysis']['variance_coefficient']
                org_details.append(f"Balanced sections (variance={variance:.3f} - AI signature)")

            if org_analysis.uniform_paragraphs:
                variance = org_analysis.details['paragraph_analysis']['variance_coefficient']
                org_details.append(f"Uniform paragraphs (variance={variance:.3f} - AI signature)")

            context_applied.append(
                f"AI-specific organizational patterns detected (NOT subject to cognitive protection): {'; '.join(org_details)}"
            )

        # Tier 1 integrity signal: gibberish / incoherence gate
        gibberish_detected = False
        gibberish_reason = ""
        gibberish_detail = ""
        if gibberish_result and gibberish_result.is_gibberish:
            gibberish_detected = True
            gibberish_reason = gibberish_result.reason
            gibberish_detail = gibberish_result.detail
            marker_counts['gibberish'] = 1
            markers_found['gibberish'] = [
                f"Incoherence detected ({gibberish_result.reason}): "
                f"{gibberish_result.detail}"
            ]
            context_applied.append(
                f"⚠ INTEGRITY: Submission flagged as incoherent — "
                f"{gibberish_result.reason}: {gibberish_result.detail}"
            )

        # Tier 1 integrity signal: unicode manipulation
        unicode_manipulation = unicode_manipulation_count > 0
        if unicode_manipulation:
            marker_counts['unicode_manipulation'] = unicode_manipulation_count
            markers_found['unicode_manipulation'] = [
                f"Text manipulation detected: {'; '.join(unicode_manipulation_details)}"
            ]
            context_applied.append(
                f"⚠ INTEGRITY: {unicode_manipulation_count} invisible/zero-width "
                f"characters detected and stripped before analysis"
            )

        # Smoking gun overrides concern level — chatbot paste artifacts are definitive
        if smoking_gun:
            concern_level = 'high'
            adjusted_concern = 'high'
            for k, v in artifact_counts.items():
                if v > 0:
                    marker_counts[f'sg_{k}'] = v
            context_applied.append(
                "⚠ SMOKING GUN: Raw chatbot artifact(s) detected — concern level forced to HIGH"
            )

        # ── Engagement reframe: compute engagement signals ──────────────
        # Maps existing HPD categories to the engagement frame from the spec.
        # These measure HOW the student engaged, not WHETHER they cheated.
        engagement_signals = self._compute_engagement_signals(
            human_presence_result=human_presence_result,
            org_analysis=org_analysis,
            specific_citation_count=specific_citation_count,
            generic_reading_ref_count=generic_reading_ref_count,
            word_count=word_count,
        )

        return AnalysisResult(
            student_id=student_id,
            student_name=student_name,
            text=text,
            word_count=word_count,
            suspicious_score=round(suspicious_score, 2),
            authenticity_score=round(authenticity_score, 2),
            marker_counts=marker_counts,
            markers_found=markers_found,
            concern_level=concern_level,
            context_adjustments_applied=context_applied,
            adjusted_suspicious_score=round(adjusted_suspicious, 2),
            adjusted_concern_level=adjusted_concern,
            organizational_analysis=org_analysis.details if org_analysis else None,
            ai_organizational_score=ai_org_score,
            human_presence_confidence=human_confidence,
            human_presence_level=human_level,
            human_presence_details=human_presence_result.__dict__ if human_presence_result else None,
            student_context=student_context,
            student_context_applied=student_context_applied,
            student_context_adjustments=student_context_adjustments,
            unicode_manipulation=unicode_manipulation,
            unicode_manipulation_details=unicode_manipulation_details,
            gibberish_detected=gibberish_detected,
            gibberish_reason=gibberish_reason,
            gibberish_detail=gibberish_detail,
            citation_count=len(citations),
            specific_citation_count=specific_citation_count,
            generic_reading_ref_count=generic_reading_ref_count,
            citations=citations,
            engagement_signals=engagement_signals,
            conversation_starters=conversation_starters,
            revision_guidance=revision_guidance,
            verification_questions=verification_questions,
            smoking_gun=smoking_gun,
            smoking_gun_details=smoking_gun_details,
            education_level=(
                self.composed_weights.education_level
                if self.composed_weights is not None else None
            ),
            population_settings=(
                self.composed_weights.population.to_dict()
                if self.composed_weights is not None else None
            ),
        )
    
    def _compute_engagement_signals(
        self,
        human_presence_result,
        org_analysis,
        specific_citation_count: int,
        generic_reading_ref_count: int,
        word_count: int,
    ) -> Dict[str, Any]:
        """Compute engagement signals from HPD categories and structural analysis.

        Maps existing analysis to the engagement frame from the spec:
          - Personal connection  ← HPD authentic_voice (15%)
          - Intellectual work    ← HPD cognitive_struggle (20%)
          - Course engagement    ← HPD contextual_grounding (35%) + citations
          - Personal investment  ← HPD emotional_stakes (20%)
          - Real-time thinking   ← HPD productive_messiness (10%)
          - Source depth         ← citation checker (specific vs generic)

        Returns a dict with per-dimension scores and an overall engagement depth.
        """
        signals: Dict[str, Any] = {}

        # Extract HPD category scores if available
        if human_presence_result is not None:
            hp = human_presence_result

            def _dim(cat_score, label: str) -> Dict[str, Any]:
                """Build one engagement dimension from an HPD CategoryScore."""
                raw = cat_score.raw_score
                weighted = cat_score.weighted_score
                markers = cat_score.marker_count
                # Normalize to 0-10 scale for display
                # HPD raw_score varies by category; use weighted for consistency
                # weighted is already relative to category weight (0-1 weight × raw)
                return {
                    'label': label,
                    'score': round(weighted, 2),
                    'marker_count': markers,
                    'markers': cat_score.markers_found[:5],  # top 5 for brevity
                }

            signals['personal_connection'] = _dim(
                hp.authentic_voice, 'Personal connection'
            )
            signals['intellectual_work'] = _dim(
                hp.cognitive_struggle, 'Intellectual work'
            )
            signals['course_engagement'] = _dim(
                hp.contextual_grounding, 'Course engagement'
            )
            signals['personal_investment'] = _dim(
                hp.emotional_stakes, 'Personal investment'
            )
            signals['real_time_thinking'] = _dim(
                hp.productive_messiness, 'Real-time thinking'
            )

            # Overall engagement depth based on total HP confidence
            hp_pct = hp.confidence_percentage
            if hp_pct >= 30.0:
                depth = 'strong'
            elif hp_pct >= 15.0:
                depth = 'moderate'
            elif hp_pct >= 5.0:
                depth = 'limited'
            else:
                depth = 'minimal'
            signals['engagement_depth'] = depth
            signals['engagement_confidence'] = round(hp_pct, 1)
        else:
            signals['engagement_depth'] = 'unavailable'
            signals['engagement_confidence'] = None

        # Source depth: specific citations vs generic reading references
        signals['source_depth'] = {
            'specific_citations': specific_citation_count,
            'generic_references': generic_reading_ref_count,
            'total': specific_citation_count + generic_reading_ref_count,
            'interpretation': (
                'Strong source engagement'
                if specific_citation_count >= 2
                else 'Some source engagement'
                if specific_citation_count >= 1 or generic_reading_ref_count >= 1
                else 'No source references detected'
            ),
        }

        # Structural note from org analysis (supplementary context, NOT verdict)
        if org_analysis:
            sent = org_analysis.details.get('sentence_analysis', {})
            structural_notes = []
            vc = sent.get('variance_coefficient', 1.0)
            if vc < 0.35:
                structural_notes.append(
                    f"Uniform sentence rhythm (VC={vc:.2f}) — common in "
                    f"ChatGPT output but also appears in careful human writers"
                )
            sd = sent.get('starter_diversity', 0.0)
            if sd >= 0.95:
                structural_notes.append(
                    f"Perfect sentence-starter diversity ({sd:.2f}) — "
                    f"AI avoids repeated first words due to architectural constraints"
                )
            cd = sent.get('comma_density', 0.0)
            if cd >= 4.5:
                structural_notes.append(
                    f"High comma density ({cd:.1f}/100 words) — "
                    f"may indicate complex AI-style subordinate clauses"
                )
            signals['structural_notes'] = structural_notes

        # Conversation opportunity: if engagement is limited/minimal,
        # suggest a conversation starter based on which dimensions are lowest
        if signals.get('engagement_depth') in ('limited', 'minimal'):
            signals['conversation_opportunity'] = True
        else:
            signals['conversation_opportunity'] = False

        return signals

    def _detect_raw_ai_artifacts(self, raw_text: str) -> tuple:
        """
        Detect copy-paste artifacts from chatbot output before HTML cleaning destroys them.

        Three scenarios:
          1. Structural HTML from AI generation (h2/h3 headers, bullet lists in prose)
          2. HTML-as-text (student pasted raw HTML source — &lt;div&gt; etc. appear as content)
          3. Unprocessed markdown in text (** bold **, ## headers, - bullet lists)

        Returns:
            (smoking_gun: bool, details: List[str], artifact_counts: Dict[str, int])
        """
        if not raw_text:
            return False, [], {}

        details = []
        counts: dict = {}

        # ── Scenario 1: Structural HTML indicating AI-generated document structure ──
        # Normal Canvas submissions have <p> tags. AI pastes add headers and lists.
        h_tags       = re.findall(r'<h[1-3][^>]*>', raw_text, re.I)
        ul_ol_tags   = re.findall(r'<[uo]l[^>]*>',  raw_text, re.I)
        li_tags      = re.findall(r'<li[^>]*>',      raw_text, re.I)
        strong_hdrs  = re.findall(
            r'<strong[^>]*>[^<]{3,60}</strong>\s*(?:<br\s*/?>|</p>)',
            raw_text, re.I
        )  # <strong>Section Title</strong> used as pseudo-headers

        counts['html_headers']    = len(h_tags)
        counts['html_lists']      = len(ul_ol_tags)
        counts['html_list_items'] = len(li_tags)
        counts['strong_headers']  = len(strong_hdrs)

        if len(h_tags) >= 2:
            details.append(
                f"Structural HTML headers detected ({len(h_tags)} <h1>/<h2>/<h3> tags) — "
                f"indicates document generated outside Canvas"
            )
        elif len(h_tags) == 1 and (ul_ol_tags or strong_hdrs):
            details.append(
                f"AI document structure: 1 HTML header + list/strong-header combo"
            )
        # Note: <ul>/<ol> with <li> items are produced by Canvas's own rich
        # text editor, so their presence alone is NOT a smoking gun.  They
        # only matter when combined with other structural artifacts (headers,
        # strong-headers) — those combos are already caught above.
        if len(strong_hdrs) >= 2:
            details.append(
                f"<strong> tags used as section headers ({len(strong_hdrs)} instances) — "
                f"structural pattern typical of AI output"
            )

        # ── Scenario 2: HTML-as-text (double-encoded / raw HTML source pasted) ──
        # When a student pastes HTML source code, Canvas stores it as
        # &lt;div&gt; etc. in the submission body.
        encoded_tags = re.findall(r'&lt;/?[a-z]{1,10}(?:\s[^&]{0,30})?&gt;', raw_text, re.I)
        counts['encoded_html_tags'] = len(encoded_tags)

        if encoded_tags:
            samples = list({t.lower() for t in encoded_tags[:6]})
            details.append(
                f"HTML code pasted as literal text ({len(encoded_tags)} encoded tags: "
                f"{', '.join(samples)}) — student likely copied raw HTML source from chatbot"
            )

        # ── Scenario 3: Unprocessed markdown in plain text ──
        # Detect after partial unescape (doesn't need full clean)
        unescaped = unescape(raw_text)
        # Strip HTML so we're only looking at text content
        text_only = re.sub(r'<[^>]+>', ' ', unescaped)

        bold_md   = re.findall(r'\*\*[^*\n]{3,80}\*\*', text_only)
        h_md      = re.findall(r'(?:^|\n)#{1,3}\s+\S', text_only)
        bullet_md = re.findall(r'(?:^|\n)\s*[-*]\s+\w', text_only)

        counts['markdown_bold']    = len(bold_md)
        counts['markdown_headers'] = len(h_md)
        counts['markdown_bullets'] = len(bullet_md)

        if len(bold_md) >= 3:
            details.append(
                f"Unprocessed markdown bold syntax ({len(bold_md)} **text** instances) — "
                f"markdown not rendered; likely pasted from a chatbot text response"
            )
        if h_md:
            details.append(
                f"Markdown headers in submission ({len(h_md)} ## / # lines) — "
                f"indicates text pasted directly from chatbot output"
            )
        if len(bullet_md) >= 3:
            details.append(
                f"Markdown bullet list ({len(bullet_md)} - / * items) in prose assignment"
            )

        smoking_gun = len(details) > 0
        return smoking_gun, details, counts

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text for analysis."""
        if not text:
            return ""
        
        # Decode HTML entities
        text = unescape(text)
        
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', ' ', text)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def _analyze_with_builtin_patterns(self, text: str) -> Tuple[float, float, Dict[str, int], Dict[str, List[str]]]:
        """
        Analyze text using built-in patterns.

        Implements cognitive diversity protection:
        - Detects neurodivergent thinking patterns
        - Applies 50% reduction to organizational bias markers when cognitive diversity markers present
        """
        text_lower = text.lower()

        suspicious_score = 0.0
        authenticity_score = 0.0
        marker_counts = {}
        markers_found = {}

        # FIRST: Check for cognitive diversity markers (protective indicators)
        cognitive_matches = []
        for phrase in COGNITIVE_DIVERSITY_MARKERS:
            count = text_lower.count(phrase.lower())
            if count > 0:
                cognitive_matches.extend([phrase] * count)
        marker_counts['cognitive_diversity'] = len(cognitive_matches)
        if cognitive_matches:
            markers_found['cognitive_diversity'] = cognitive_matches[:5]

        # Determine protection level
        # Uses self._cognitive_protection_floor (0.5 standard; 0.3 with ND-aware overlay)
        cognitive_protection_multiplier = 1.0
        if len(cognitive_matches) >= 2:
            if len(cognitive_matches) >= 4:
                cognitive_protection_multiplier = self._cognitive_protection_floor  # Strong protection
            else:
                cognitive_protection_multiplier = 0.7  # Moderate: 30% reduction

        # Cognitive diversity markers boost authenticity
        # Weight from _marker_weights (boosted by ND-aware overlay to 0.9 when on)
        cog_div_weight = self._marker_weights.get('cognitive_diversity', 0.6)
        authenticity_score += min(len(cognitive_matches) * cog_div_weight, 9.0)

        # Check AI transitions (ORGANIZATIONAL BIAS - subject to protection)
        # Notes mode: smooth transitions are EXTRA suspicious (amplified 2×)
        transition_matches = []
        for phrase in AI_TRANSITIONS:
            count = text_lower.count(phrase.lower())
            if count > 0:
                transition_matches.extend([phrase] * count)
        marker_counts['ai_transitions'] = len(transition_matches)
        if transition_matches:
            markers_found['ai_transitions'] = transition_matches[:5]
        transition_amplifier = 2.0 if self._invert_sentence_signals else 1.0
        suspicious_score += len(transition_matches) * self._marker_weights.get('ai_transitions', 0.5) * cognitive_protection_multiplier * transition_amplifier

        # Clustering bonus: multiple AI markers in short text is very suspicious
        # Also subject to cognitive protection
        if len(transition_matches) >= 3 and len(text.split()) < 500:
            suspicious_score += 2.0 * cognitive_protection_multiplier

        # Check generic phrases (ORGANIZATIONAL BIAS - subject to protection)
        # Notes mode: generic essay-summary phrases are EXTRA suspicious (amplified 2×)
        generic_matches = []
        for phrase in GENERIC_PHRASES:
            count = text_lower.count(phrase.lower())
            if count > 0:
                generic_matches.extend([phrase] * count)
        marker_counts['generic_phrases'] = len(generic_matches)
        if generic_matches:
            markers_found['generic_phrases'] = generic_matches[:5]
        generic_amplifier = 2.0 if self._invert_sentence_signals else 1.0
        suspicious_score += len(generic_matches) * self._marker_weights.get('generic_phrases', 0.4) * cognitive_protection_multiplier * generic_amplifier

        # Check inflated vocabulary (ORGANIZATIONAL BIAS - subject to protection)
        inflated_matches = []
        for inflated, simple in INFLATED_VOCAB:
            count = text_lower.count(inflated.lower())
            if count > 0:
                inflated_matches.extend([inflated] * count)
        marker_counts['inflated_vocabulary'] = len(inflated_matches)
        if inflated_matches:
            markers_found['inflated_vocabulary'] = inflated_matches[:5]
        # Apply cognitive protection (weight from _marker_weights)
        suspicious_score += len(inflated_matches) * self._marker_weights.get('inflated_vocabulary', 0.3) * cognitive_protection_multiplier

        # Check personal markers (presence is GOOD — unless personal_voice_authentic=False)
        # personal_voice_authentic=False: notes/essays where first-person ≠ authenticity signal
        personal_count = 0
        personal_matches = []
        for pattern in PERSONAL_MARKERS:
            matches = re.findall(pattern, text, re.IGNORECASE)
            personal_count += len(matches)
            if matches:
                personal_matches.extend([m.strip() for m in matches[:3]])
        marker_counts['personal_voice'] = personal_count
        if personal_matches:
            markers_found['personal_voice'] = personal_matches[:5]
        if self._personal_voice_authentic:
            authenticity_score += min(personal_count * self._marker_weights.get('personal_voice', 0.5), 5.0)

        # Check emotional markers (presence is GOOD — unless personal_voice_authentic=False)
        emotional_count = 0
        emotional_matches = []
        for phrase in EMOTIONAL_MARKERS:
            count = text_lower.count(phrase.lower())
            emotional_count += count
            if count > 0:
                emotional_matches.append(phrase)
        marker_counts['emotional_language'] = emotional_count
        if emotional_matches:
            markers_found['emotional_language'] = emotional_matches[:5]
        if self._personal_voice_authentic:
            authenticity_score += min(emotional_count * self._marker_weights.get('emotional_language', 0.8), 4.0)

        # Check for specific details (proper nouns are GOOD)
        proper_nouns = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
        # Filter out sentence starters
        specific_count = len([n for n in proper_nouns if len(n) > 3])
        marker_counts['specific_details'] = specific_count
        authenticity_score += min(specific_count * 0.2, 3.0)

        return suspicious_score, authenticity_score, marker_counts, markers_found, cognitive_protection_multiplier
    
    def _apply_profile_weights(self, score: float, marker_counts: Dict[str, int]) -> float:
        """Apply profile-specific weight multipliers to SUSPICIOUS markers only."""
        multipliers = self.profile.get('weight_multipliers', {})
        
        # Only count suspicious markers, not authenticity markers
        suspicious_markers = {'ai_transitions', 'generic_phrases', 'inflated_vocabulary'}
        
        adjusted_score = 0.0
        for marker_id, count in marker_counts.items():
            if marker_id not in suspicious_markers:
                continue
            base_weight = 0.3  # Base weight per marker
            multiplier = multipliers.get(marker_id, 1.0)
            adjusted_score += count * base_weight * multiplier
        
        return adjusted_score
    
    # Minimum word count for reliable AIC analysis.  Below this threshold
    # there simply isn't enough text for authentic-voice markers to appear,
    # so low authenticity scores are expected and not evidence of AI use.
    MIN_WORDS_FOR_RELIABLE_ANALYSIS = 30

    def _determine_concern_level(
        self, suspicious: float, authenticity: float,
        word_count: int | None = None,
    ) -> str:
        """Determine concern level based on scores.

        Authenticity thresholds are mode-sensitive: when the assignment mode
        suppresses personal voice (e.g. essay, lab, outline), the expected
        authenticity baseline is lower and thresholds shift accordingly.
        Without this adjustment, formal assignment modes produce false
        positives because students writing well-structured essays don't
        use personal anecdotes — which is correct behavior, not a signal
        of disengagement.

        If *word_count* is provided and below
        ``MIN_WORDS_FOR_RELIABLE_ANALYSIS``, the result is capped at ``'low'``
        because very short submissions lack enough text for authenticity
        markers to register reliably.
        """
        # Mode-sensitive authenticity thresholds:
        # weight_personal_voice < 0.8 means the mode expects less personal voice,
        # so lower the authenticity thresholds proportionally.
        wpv = 1.0
        if self._aic_config is not None:
            wpv = self._aic_config.get('weight_personal_voice', 1.0)
        # Scale: wpv=1.0 → no shift; wpv=0.5 → thresholds halved
        auth_scale = max(0.3, min(1.0, wpv))

        auth_high = 2.0 * auth_scale       # default 2.0, essay mode ~1.0
        auth_elevated = 1.5 * auth_scale    # default 1.5, essay mode ~0.75
        auth_moderate = 3.0 * auth_scale    # default 3.0, essay mode ~1.5

        # --- score-based determination (mode-adaptive) ---
        # High: High suspicious AND low authenticity
        if suspicious > 5.0 and authenticity < auth_high:
            level = 'high'
        # Elevated: Moderate suspicious OR very low authenticity
        elif suspicious > 4.0 or authenticity < auth_elevated:
            level = 'elevated'
        # Moderate: Some concerning patterns
        elif suspicious > 2.5 or authenticity < auth_moderate:
            level = 'moderate'
        # Low: Minor issues
        elif suspicious > 1.5:
            level = 'low'
        else:
            level = 'none'

        # --- brevity cap ------------------------------------------------
        # Very short submissions don't contain enough text for human-voice
        # markers to surface.  Flagging them would penalise brevity, not
        # detect AI use.
        if (
            word_count is not None
            and word_count < self.MIN_WORDS_FOR_RELIABLE_ANALYSIS
            and level not in ('low', 'none')
        ):
            level = 'low'

        return level
    
    def _get_conversation_starters(self, concern_level: str) -> List[str]:
        """Get appropriate conversation starters."""
        starters = {
            'high': [
                "Can you walk me through how you approached this assignment?",
                "I'd love to hear more about your thinking process.",
                "Can you tell me about a specific moment that influenced your writing?",
                "What was the most challenging part of this for you?"
            ],
            'elevated': [
                "Can you expand on some of the ideas in your paper?",
                "How did you decide on this approach?",
                "What specific examples from your experience informed this?"
            ],
            'moderate': [
                "Can you add some more personal examples?",
                "I'd like to hear more of your own voice in this."
            ]
        }
        return starters.get(concern_level, [])
    
    def _get_revision_guidance(self, concern_level: str) -> List[str]:
        """Get revision guidance."""
        if concern_level not in ['high', 'elevated']:
            return []
        
        if self.profile_id == 'personal_reflection':
            return [
                "Add 2-3 specific sensory details from your actual experience",
                "Include at least one specific person, place, or time from your life",
                "Remove generic statements that could apply to anyone"
            ]
        else:
            return [
                "Add more specific examples and details",
                "Include your own analysis and interpretation",
                "Make sure your own thinking is clearly present"
            ]
    
    def _get_verification_questions(self, concern_level: str) -> List[str]:
        """Get verification questions for instructor."""
        if concern_level not in ['high', 'elevated']:
            return []
        
        return [
            "What specific moment or experience prompted this reflection?",
            "Can you describe in more detail what you actually saw/heard/felt?",
            "How did your understanding change as you wrote this?",
            "Which part of the reading specifically connected to your experience?"
        ]

    def _process_student_context(self, student_context: str, suspicious_score: float,
                                 ai_org_score: float) -> tuple[float, float, List[str]]:
        """
        Process optional student-provided context about their writing process.

        PHASE 6: Privacy-focused optional field. NO DIAGNOSIS REQUIRED.

        Args:
            student_context: Student's description of writing process
            suspicious_score: Current suspicious score
            ai_org_score: AI organizational score

        Returns:
            (adjusted_suspicious, adjusted_org, adjustments_list)
        """
        if not student_context:
            return suspicious_score, ai_org_score, []

        adjustments = []
        context_lower = student_context.lower()

        # Track adjustments
        suspicious_adjustment = 1.0
        org_adjustment = 1.0

        # Non-linear / organizational differences
        non_linear_indicators = [
            'non-linear', 'nonlinear', 'non linear',
            'jump around', 'all over the place',
            'scattered', 'organize differently',
            'neurodivergent', 'adhd', 'autism', 'autistic',
            'executive function', 'processing disorder'
        ]

        if any(indicator in context_lower for indicator in non_linear_indicators):
            # Reduce weight on organizational patterns
            org_adjustment = 0.5  # 50% reduction
            adjustments.append(
                "Student reports non-linear writing process - organizational patterns weighted 50% less"
            )

        # ESL / Language learning
        esl_indicators = [
            'second language', 'third language', 'esl', 'ell',
            'english is not my first', 'learning english',
            'non-native', 'nonnative', 'foreign language'
        ]

        if any(indicator in context_lower for indicator in esl_indicators):
            # Focus more on engagement than polish
            suspicious_adjustment = 0.7  # 30% reduction
            adjustments.append(
                "Student reports English as additional language - focusing on engagement over polish (30% reduction)"
            )

        # Assistive technology use
        assistive_tech_indicators = [
            'voice to text', 'voice-to-text', 'speech to text',
            'dictation', 'screen reader', 'assistive technology',
            'accessibility tool', 'grammarly', 'grammar checker'
        ]

        if any(indicator in context_lower for indicator in assistive_tech_indicators):
            adjustments.append(
                "Student reports assistive technology use - polished output expected and legitimate"
            )
            # Don't reduce scores, just note for instructor awareness

        # Writing process descriptions (metacognitive awareness = human)
        process_indicators = [
            'i usually', 'i tend to', 'my process',
            'i write best', 'i draft', 'i revise',
            'i struggle with', 'it helps me to'
        ]

        if any(indicator in context_lower for indicator in process_indicators):
            adjustments.append(
                "Student demonstrates metacognitive awareness of writing process - increases authenticity"
            )

        # Apply adjustments
        adjusted_suspicious = suspicious_score * suspicious_adjustment
        adjusted_org = ai_org_score * org_adjustment

        return adjusted_suspicious, adjusted_org, adjustments


# =============================================================================
# PEER COMPARISON
# =============================================================================

class PeerComparisonAnalyzer:
    """
    Statistical outlier detection based on peer comparison.
    
    Philosophy: A score that's unusual for THIS class is more meaningful
    than hitting an absolute threshold. Adapts to each cohort's patterns.
    """
    
    def __init__(self, outlier_percentile: float = 90.0):
        """
        Initialize peer comparison analyzer.
        
        Args:
            outlier_percentile: Percentile above which to flag as outlier
                               (95 for community college, 90 for others)
        """
        self.outlier_percentile = outlier_percentile
    
    def analyze_cohort(self, results: List[AnalysisResult]) -> Dict[str, Any]:
        """
        Analyze a cohort of submissions and identify statistical outliers.
        
        Returns summary statistics and updates each result with percentile info.
        """
        if len(results) < 3:
            return {"error": "Need at least 3 submissions for peer comparison"}
        
        # Extract scores
        suspicious_scores = [r.suspicious_score for r in results]
        authenticity_scores = [r.authenticity_score for r in results]
        
        # Calculate statistics
        sus_mean = statistics.mean(suspicious_scores)
        sus_stdev = statistics.stdev(suspicious_scores) if len(suspicious_scores) > 1 else 0
        auth_mean = statistics.mean(authenticity_scores)
        auth_stdev = statistics.stdev(authenticity_scores) if len(authenticity_scores) > 1 else 0
        
        # Calculate percentile threshold
        sorted_sus = sorted(suspicious_scores)
        threshold_idx = int(len(sorted_sus) * self.outlier_percentile / 100)
        sus_threshold = sorted_sus[min(threshold_idx, len(sorted_sus) - 1)]
        
        # Update each result with percentile and outlier info
        outliers = []
        for result in results:
            # Calculate percentile
            below_count = sum(1 for s in suspicious_scores if s < result.suspicious_score)
            result.suspicious_percentile = round(100 * below_count / len(suspicious_scores), 1)
            
            below_count = sum(1 for s in authenticity_scores if s < result.authenticity_score)
            result.authenticity_percentile = round(100 * below_count / len(authenticity_scores), 1)
            
            # Check if outlier
            outlier_reasons = []
            
            if result.suspicious_score > sus_threshold:
                outlier_reasons.append(f"Suspicious score in top {100 - self.outlier_percentile:.0f}% of class")
            
            if sus_stdev > 0:
                z_score = (result.suspicious_score - sus_mean) / sus_stdev
                if z_score > 2.0:
                    outlier_reasons.append(f"Z-score {z_score:.1f} (>2 standard deviations)")
            
            if result.authenticity_score < auth_mean - (1.5 * auth_stdev) and auth_stdev > 0:
                outlier_reasons.append("Low authenticity compared to peers")
            
            if outlier_reasons:
                result.is_outlier = True
                result.outlier_reasons = outlier_reasons
                outliers.append(result)
        
        return {
            "total_submissions": len(results),
            "suspicious_mean": round(sus_mean, 2),
            "suspicious_stdev": round(sus_stdev, 2),
            "suspicious_threshold": round(sus_threshold, 2),
            "authenticity_mean": round(auth_mean, 2),
            "authenticity_stdev": round(auth_stdev, 2),
            "outlier_count": len(outliers),
            "outliers": outliers
        }


# =============================================================================
# REPORT GENERATOR
# =============================================================================

class ReportGenerator:
    """
    Generates pedagogically-informed reports.
    
    Philosophy: Reports are for INSTRUCTORS, not verdicts.
    Frame findings as conversation starters, not accusations.
    """
    
    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize report generator.

        Args:
            output_dir: Where to save reports

        Note: File cleanup is handled by the main run_autograder.py cleanup system.
              See Settings > Configure automatic cleanup for options.
        """
        self.output_dir = output_dir or (get_output_base_dir() / "Academic Dishonesty Reports")
        # Security: Restrict permissions to owner only
        self.output_dir.mkdir(parents=True, exist_ok=True, mode=0o700)
    
    def generate_report(self,
                        assignment_name: str,
                        profile_name: str,
                        results: List[AnalysisResult],
                        cohort_stats: Optional[Dict] = None,
                        context_profile: str = "community_college") -> Path:
        """Generate a text report."""
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c for c in assignment_name if c.isalnum() or c in " -_")[:50]
        filepath = self.output_dir / f"{safe_name}_{timestamp}_report.txt"
        
        lines = []
        
        # Header
        lines.append("=" * 75)
        lines.append("ACADEMIC INTEGRITY ANALYSIS REPORT")
        lines.append("=" * 75)
        lines.append("")
        lines.append(f"Assignment: {assignment_name}")
        lines.append(f"Analysis Profile: {profile_name}")
        lines.append(f"Context Profile: {context_profile}")
        lines.append(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append(f"Total Submissions: {len(results)}")
        lines.append("")
        
        # Important framing
        lines.append("-" * 75)
        lines.append("IMPORTANT: HOW TO USE THIS REPORT")
        lines.append("-" * 75)
        lines.append("""
This report identifies INDICATORS that warrant instructor review.
It does NOT provide proof of academic dishonesty.

WHAT THIS TOOL IS:
  ✓ A conversation starter for discussing student work
  ✓ An outlier detector (peer comparison, not absolute judgment)
  ✓ A pedagogical support for upholding learning objectives

WHAT THIS TOOL IS NOT:
  ✗ A verdict or proof of cheating
  ✗ An AI detector (detects dishonest USE, not all AI use)
  ✗ Infallible (false positives occur - especially for ESL students)

RECOMMENDED APPROACH:
  1. Review flagged submissions yourself before any conversation
  2. Use conversation starters provided - don't accuse
  3. Consider student's other work for context
  4. Frame as learning opportunity when discussing with students
""")
        lines.append("")
        
        # Cohort summary
        if cohort_stats:
            lines.append("-" * 75)
            lines.append("COHORT STATISTICAL SUMMARY")
            lines.append("-" * 75)
            lines.append(f"  Suspicious Score - Mean: {cohort_stats.get('suspicious_mean', 'N/A')}, "
                        f"StDev: {cohort_stats.get('suspicious_stdev', 'N/A')}")
            lines.append(f"  Authenticity Score - Mean: {cohort_stats.get('authenticity_mean', 'N/A')}, "
                        f"StDev: {cohort_stats.get('authenticity_stdev', 'N/A')}")
            lines.append(f"  Outliers Identified: {cohort_stats.get('outlier_count', 0)}")
            lines.append("")
        
        # Summary by concern level
        lines.append("-" * 75)
        lines.append("SUMMARY BY CONCERN LEVEL")
        lines.append("-" * 75)
        
        high_concern = [r for r in results if r.concern_level == 'high']
        elevated = [r for r in results if r.concern_level == 'elevated']
        moderate = [r for r in results if r.concern_level == 'moderate']
        low = [r for r in results if r.concern_level == 'low']
        clean = [r for r in results if r.concern_level == 'none']
        smoking_guns = [r for r in results if r.smoking_gun]

        if smoking_guns:
            lines.append(f"  !! SMOKING GUNS:  {len(smoking_guns):3d} (raw chatbot paste artifacts detected!)")
        lines.append(f"  HIGH CONCERN:     {len(high_concern):3d} (recommend structured conversation)")
        lines.append(f"  ELEVATED:         {len(elevated):3d} (recommend brief check-in)")
        lines.append(f"  MODERATE:         {len(moderate):3d} (note for pattern tracking)")
        lines.append(f"  LOW:              {len(low):3d} (feedback only)")
        lines.append(f"  NO CONCERNS:      {len(clean):3d}")
        lines.append("")
        
        # Detailed reports for high/elevated concern
        if high_concern or elevated:
            lines.append("=" * 75)
            lines.append("DETAILED ANALYSIS: HIGH & ELEVATED CONCERN")
            lines.append("=" * 75)
            
            for result in high_concern + elevated:
                lines.extend(self._format_detailed_result(result))
                lines.append("")
        
        # Brief listing for moderate concern
        if moderate:
            lines.append("-" * 75)
            lines.append("MODERATE CONCERN (Brief Summary)")
            lines.append("-" * 75)
            for result in moderate:
                emoji = "🟡"
                lines.append(f"  {emoji} {result.student_name}: suspicious={result.suspicious_score:.1f}, "
                           f"authenticity={result.authenticity_score:.1f}")
            lines.append("")
        
        # Footer
        lines.append("=" * 75)
        lines.append("END OF REPORT")
        lines.append("=" * 75)
        lines.append("")
        lines.append("Remember: All flags require human judgment. When in doubt,")
        lines.append("have a conversation with the student before drawing conclusions.")
        
        # Write file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
        
        return filepath

    def _format_detailed_result(self, result: AnalysisResult) -> List[str]:
        """Format a single detailed result."""
        lines = []
        
        concern_emoji = {
            'high': '🔴',
            'elevated': '🟠',
            'moderate': '🟡',
            'low': '🟢',
            'none': '⚪'
        }
        
        emoji = concern_emoji.get(result.concern_level, '⚪')
        
        lines.append("-" * 50)

        # Smoking gun banner — shown before everything else
        if result.smoking_gun:
            lines.append("  !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            lines.append("  !! SMOKING GUN: CHATBOT PASTE DETECTED !!")
            lines.append("  !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            for detail in result.smoking_gun_details:
                lines.append(f"    >> {detail}")
            lines.append("  Note: Raw chatbot artifacts were found in the submission.")
            lines.append("  This is distinct from 'AI-assisted writing' — the student")
            lines.append("  literally copied output from a chatbot interface.")
            lines.append("")

        lines.append(f"{emoji} {result.student_name}")
        lines.append(f"   Concern Level: {result.concern_level.upper()}")
        if result.smoking_gun:
            lines.append(f"   (Forced HIGH due to smoking gun artifact detection)")
        lines.append(f"   Word Count: {result.word_count}")
        lines.append("")

        # Human presence confidence
        if result.human_presence_confidence is not None:
            lines.append(f"   HUMAN PRESENCE: {result.human_presence_confidence:.0f}% "
                         f"({result.human_presence_level or 'unknown'})")
            lines.append("")

        lines.append("   SCORES:")
        lines.append(f"     Suspicious Markers: {result.suspicious_score:.2f}")
        if result.suspicious_percentile is not None:
            lines[-1] += f" ({result.suspicious_percentile:.0f}th percentile)"
        lines.append(f"     Authenticity Markers: {result.authenticity_score:.2f}")
        if result.authenticity_percentile is not None:
            lines[-1] += f" ({result.authenticity_percentile:.0f}th percentile)"
        lines.append("")
        
        if result.is_outlier and result.outlier_reasons:
            lines.append("   STATISTICAL OUTLIER:")
            for reason in result.outlier_reasons:
                lines.append(f"     → {reason}")
            lines.append("")
        
        if result.context_adjustments_applied:
            lines.append("   CONTEXT ADJUSTMENTS APPLIED:")
            for ctx in result.context_adjustments_applied:
                lines.append(f"     • {ctx}")
            lines.append("")
        
        if result.markers_found:
            lines.append("   KEY MARKERS DETECTED:")
            for marker_type, instances in list(result.markers_found.items())[:5]:
                lines.append(f"     {marker_type}: {len(instances)} instance(s)")
                # Privacy: Do not include actual student text in reports
                # Quantitative data (counts) is sufficient for instructor review
            lines.append("")
        
        if result.conversation_starters:
            lines.append("   CONVERSATION STARTERS:")
            for starter in result.conversation_starters[:3]:
                lines.append(f"     • \"{starter}\"")
            lines.append("")
        
        if result.verification_questions:
            lines.append("   VERIFICATION QUESTIONS:")
            for q in result.verification_questions[:3]:
                lines.append(f"     • {q}")
        
        return lines


# =============================================================================
# CANVAS INTEGRATION
# =============================================================================

def get_courses():
    """Fetch list of courses for the current user."""
    if not HAS_REQUESTS or not API_TOKEN:
        print("Error: Canvas API not configured.")
        return []
    
    url = f"{CANVAS_BASE_URL}/api/v1/courses"
    params = {"enrollment_state": "active", "per_page": 100}
    
    try:
        response = requests.get(url, headers=HEADERS, params=params)
        response.raise_for_status()
        return response.json()
    except Exception as e:
        print(f"Error fetching courses: {e}")
        return []


def get_assignments(course_id: int):
    """Fetch assignments for a course."""
    if not HAS_REQUESTS or not API_TOKEN:
        return []
    
    url = f"{CANVAS_BASE_URL}/api/v1/courses/{course_id}/assignments"
    params = {"per_page": 100}
    
    try:
        response = requests.get(url, headers=HEADERS, params=params)
        response.raise_for_status()
        return response.json()
    except Exception as e:
        print(f"Error fetching assignments: {e}")
        return []


def _extract_attachment_text(attachment: Dict[str, Any]) -> str:
    """
    Download a submission attachment and extract its text content.
    Supports .txt, .md, .html, .docx, .doc, .odt, .pdf, .rtf files.
    Returns empty string on failure or unsupported file type.
    """
    url = attachment.get("url")
    if not url:
        return ""
    filename = (attachment.get("filename") or "").lower()
    content_type = (attachment.get("content-type") or "").lower()
    try:
        resp = requests.get(url, headers=HEADERS, timeout=60)
        resp.raise_for_status()
        data = resp.content

        # Plain text / Markdown
        if filename.endswith((".txt", ".md")) or content_type.startswith("text/plain"):
            return data.decode("utf-8", errors="replace")

        # HTML — strip tags
        if filename.endswith((".html", ".htm")) or "text/html" in content_type:
            import re
            return re.sub(r"<[^>]+>", " ", data.decode("utf-8", errors="replace")).strip()

        # Word document (.docx)
        if filename.endswith(".docx") or "openxmlformats-officedocument.wordprocessingml" in content_type:
            import io
            from docx import Document
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # Old binary Word (.doc) — try python-docx as best-effort
        if filename.endswith(".doc") or content_type == "application/msword":
            try:
                import io
                from docx import Document
                doc = Document(io.BytesIO(data))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if text:
                    return text
            except Exception:
                pass
            return ""

        # OpenDocument Text (.odt)
        if filename.endswith(".odt") or "application/vnd.oasis.opendocument.text" in content_type:
            import io
            from odf.opendocument import load as odf_load
            from odf.text import P
            from odf import teletype
            doc = odf_load(io.BytesIO(data))
            return "\n".join(
                teletype.extractText(p)
                for p in doc.getElementsByType(P)
                if teletype.extractText(p).strip()
            )

        # PDF — extract raw text
        if filename.endswith(".pdf") or content_type.startswith("application/pdf"):
            import io
            from pdfminer.high_level import extract_text as _pdf_extract
            return _pdf_extract(io.BytesIO(data)) or ""

        # RTF
        if filename.endswith(".rtf") or content_type in ("application/rtf", "text/rtf"):
            from striprtf.striprtf import rtf_to_text
            return rtf_to_text(data.decode("utf-8", errors="replace"))

    except ImportError:
        pass  # dependency not installed — degrade gracefully
    except Exception:
        pass
    return ""


def _get_discussion_submissions(course_id: int, assignment_id: int):
    """Fetch discussion entries as pseudo-submissions for AIC analysis.

    For discussion-type assignments, Canvas stores text in discussion entries
    rather than submission bodies.  This function detects discussion assignments,
    fetches entries via the /discussion_topics/{id}/view endpoint, and returns
    them in the same shape as regular submissions so the caller needs no changes.

    Returns ``None`` when the assignment is *not* a discussion — the caller
    should fall back to normal ``get_submissions``.
    """
    if not HAS_REQUESTS or not API_TOKEN:
        return None

    # ── Fetch assignment metadata to check submission_types ──────────
    try:
        url = f"{CANVAS_BASE_URL}/api/v1/courses/{course_id}/assignments/{assignment_id}"
        resp = requests.get(url, headers=HEADERS, timeout=30)
        resp.raise_for_status()
        assignment = resp.json()
    except Exception:
        return None

    submission_types = assignment.get("submission_types", [])
    if "discussion_topic" not in submission_types:
        return None  # not a discussion — caller should use get_submissions

    # ── Resolve the discussion topic ID ──────────────────────────────
    topic = assignment.get("discussion_topic")
    topic_id = topic.get("id") if isinstance(topic, dict) else None
    if not topic_id:
        topic_id = assignment.get("discussion_topic_id")
    if not topic_id:
        return None

    # ── Fetch the full discussion view (entries + nested replies) ────
    try:
        view_url = (f"{CANVAS_BASE_URL}/api/v1/courses/{course_id}"
                    f"/discussion_topics/{topic_id}/view")
        resp = requests.get(view_url, headers=HEADERS, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception:
        return None

    # ── Flatten entries and group by user_id ─────────────────────────
    raw_entries = data.get("view", [])
    user_texts: Dict[int, List[str]] = defaultdict(list)

    def _collect(entry):
        uid = entry.get("user_id")
        msg = entry.get("message") or ""
        if uid and msg.strip():
            user_texts[uid].append(msg)
        for reply in entry.get("replies", []):
            _collect(reply)

    for entry in raw_entries:
        _collect(entry)

    if not user_texts:
        return []

    # ── Build participant name lookup ────────────────────────────────
    participants = {p["id"]: p.get("display_name", f"Student {p['id']}")
                    for p in data.get("participants", [])}

    # ── Return pseudo-submission dicts (same shape the caller expects) ─
    pseudo = []
    for uid, texts in user_texts.items():
        pseudo.append({
            "user_id": uid,
            "user": {"name": participants.get(uid, f"Student {uid}")},
            "body": "\n\n".join(texts),
            "workflow_state": "submitted",
            "submitted_at": None,
        })
    return pseudo


def get_submissions(course_id: int, assignment_id: int):
    """Fetch all submissions for an assignment."""
    if not HAS_REQUESTS or not API_TOKEN:
        return []
    
    url = f"{CANVAS_BASE_URL}/api/v1/courses/{course_id}/assignments/{assignment_id}/submissions"
    params = {"per_page": 100, "include[]": ["user"]}
    
    all_submissions = []
    
    try:
        while url:
            response = requests.get(url, headers=HEADERS, params=params)
            response.raise_for_status()
            submissions = response.json()
            all_submissions.extend(submissions)
            
            # Check for pagination
            links = response.headers.get("Link", "")
            url = None
            for link in links.split(","):
                if 'rel="next"' in link:
                    url = link.split(";")[0].strip().strip("<>")
                    params = {}  # URL already contains params
                    break
        
        return all_submissions
    except Exception as e:
        print(f"Error fetching submissions: {e}")
        return []


# =============================================================================
# BATCH PROCESSING
# =============================================================================

def analyze_assignment(course_id: int,
                       assignment_id: int,
                       profile_id: str = "standard",
                       context_profile: str = "community_college",
                       assignment_type: Optional[str] = None,
                       composed_weights=None,
                       aic_config: Optional[dict] = None,
                       course_name: str = "",
                       generate_report: bool = True,
                       submissions: Optional[list] = None) -> Tuple[List[AnalysisResult], Path]:
    """
    Analyze all submissions for an assignment.

    Args:
        course_id: Canvas course ID
        assignment_id: Canvas assignment ID
        profile_id: Assignment profile (legacy)
        context_profile: Population context profile (legacy)
        assignment_type: v3.0 AIC mode ('notes', 'discussion', 'draft', 'personal',
                         'essay', 'lab'). When provided, configures signal polarity
                         and personal-voice handling for the specific assignment type.
        composed_weights: Optional ComposedWeights from WeightComposer (Phase 8).
        submissions: Optional pre-fetched (and preprocessed) submission list.
                     When provided, skips Canvas fetch — used by RunWorker so
                     preprocessing (translation, transcription) runs once.

    Returns:
        Tuple of (list of results, path to report file)
    """
    if submissions is not None:
        print(f"\nUsing {len(submissions)} pre-fetched submissions.")
    else:
        print(f"\nFetching submissions...")
        # Try discussion-entry path first (returns None for non-discussions)
        submissions = _get_discussion_submissions(course_id, assignment_id)
        if submissions is not None:
            print(f"Found {len(submissions)} discussion participants.")
        else:
            submissions = get_submissions(course_id, assignment_id)
            if not submissions:
                print("No submissions found.")
                return [], None
            print(f"Found {len(submissions)} submissions.")

    # Initialize analyzer
    # aic_config (Phase 9) takes precedence; assignment_type (Phase 7) is legacy fallback
    _at = None if aic_config else assignment_type
    analyzer = DishonestyAnalyzer(
        profile_id=profile_id,
        context_profile=context_profile,
        assignment_type=_at,
        composed_weights=composed_weights,
        aic_config=aic_config,
    )
    
    # Analyze each submission
    results = []
    errors = []
    skipped_no_text = 0
    skipped_unsubmitted = 0
    for sub in submissions:
        try:
            # Get student info
            user = sub.get("user", {})
            student_id = str(sub.get("user_id", "unknown"))
            student_name = user.get("name", f"Student {student_id}")

            # Skip unsubmitted
            ws = sub.get("workflow_state", "")
            if ws in ("unsubmitted", "deleted"):
                skipped_unsubmitted += 1
                continue

            # Get submission text — inline body first, then attached files
            body = sub.get("body", "") or ""
            if not body.strip():
                for att in (sub.get("attachments") or []):
                    body += "\n" + _extract_attachment_text(att)
                body = body.strip()

            # Skip if still no text after extraction
            if not body.strip():
                skipped_no_text += 1
                continue

            print(f"  Analyzing: {student_name}...")
            result = analyzer.analyze_text(body, student_id, student_name)
            results.append(result)
        except Exception as e:
            # Don't let one submission failure stop the entire batch
            error_msg = f"Error analyzing {student_name}: {str(e)}"
            print(f"  Warning: {error_msg}")
            print(f"  Continuing with remaining submissions...")
            errors.append(error_msg)
            continue

    if skipped_no_text:
        print(f"  Skipped {skipped_no_text} submission(s) with no extractable text"
              f" (file-only or empty)")

    if not results:
        print("No text submissions to analyze.")
        return [], None

    # Run peer comparison
    print(f"\nRunning peer comparison on {len(results)} submissions...")
    peer_analyzer = PeerComparisonAnalyzer(
        outlier_percentile=95.0 if context_profile == "community_college" else 90.0
    )
    cohort_stats = peer_analyzer.analyze_cohort(results)

    # Save results to RunStore (SQLite) so they appear in the Prior Runs dashboard
    try:
        from automation.run_store import RunStore
        store = RunStore()
        # Get assignment name for storage (reuse what we fetched for the report below)
        _assignments = get_assignments(course_id)
        _assignment_name = next(
            (a.get("name", "Unknown") for a in _assignments if a.get("id") == assignment_id),
            "Unknown Assignment",
        )
        _course_name = course_name or f"Course {course_id}"
        for result in results:
            _sub = next((s for s in submissions if str(s.get("user_id")) == str(result.student_id)), {})
            store.save_result(
                result,
                course_id=str(course_id),
                course_name=_course_name,
                assignment_id=str(assignment_id),
                assignment_name=_assignment_name,
                submitted_at=_sub.get("submitted_at"),
                context_profile=context_profile,
                submission_body=(_sub.get("body") or "").strip(),
            )
    except Exception as _e:
        print(f"  ⚠ RunStore save skipped: {_e}")

    # ── Mechanism 1: Cohort Calibration ──────────────────────────────────
    # Compute class-relative baselines from fast AIC signals and annotate
    # each student result with class-relative percentile interpretations.
    if HAS_COHORT_CALIBRATION and len(results) >= 3:
        try:
            calibrator = CohortCalibrator()

            # Extract signal vectors from all results
            signal_vectors = []
            vector_map = {}  # student_id → signal vector
            for r in results:
                vec = extract_signal_vector(r)
                if vec is not None:
                    signal_vectors.append(vec)
                    vector_map[r.student_id] = vec

            if len(signal_vectors) >= 3:
                # Compute class distributions
                distributions = calibrator.compute_class_distributions(signal_vectors)

                # Check for prior baseline (EMA evolution)
                _edu_level = (
                    composed_weights.education_level
                    if composed_weights is not None
                    else context_profile
                )
                _prior_used = False
                _prior_weight = None

                try:
                    from automation.run_store import RunStore as _RS
                    _baseline_store = _RS()
                    _prev_baseline = _baseline_store.get_class_baseline(
                        str(course_id), assignment_mode=assignment_type
                    )

                    if _prev_baseline and _prev_baseline.get('signals'):
                        # Evolve baseline with EMA
                        prev_signals = _prev_baseline['signals']
                        for sig_name, stats in distributions.items():
                            if sig_name in prev_signals:
                                prev_mean = prev_signals[sig_name].get('mean')
                                if prev_mean is not None:
                                    stats['mean'] = calibrator.evolve_baseline(
                                        stats['mean'], prev_mean
                                    )
                    else:
                        # Cold start: apply Bayesian blending with education-level priors
                        _prior_weight = max(10, 25 - len(signal_vectors))
                        _prior_used = True
                        for sig_name, stats in distributions.items():
                            stats['mean'] = calibrator.cold_start_baseline(
                                stats['mean'],
                                len(signal_vectors),
                                education_level=_edu_level or "community_college",
                                signal_name=sig_name,
                            )

                    # Save baseline to RunStore
                    _run_id = f"{course_id}_{assignment_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                    _baseline_store.save_class_baseline(
                        run_id=_run_id,
                        course_id=str(course_id),
                        assignment_mode=assignment_type,
                        education_level=_edu_level,
                        n_students=len(signal_vectors),
                        signals=distributions,
                        prior_used=_prior_used,
                        prior_weight=_prior_weight,
                    )
                except Exception as _be:
                    print(f"  ⚠ Baseline persistence skipped: {_be}")

                # Annotate each result with class-relative interpretations
                for r in results:
                    vec = vector_map.get(r.student_id)
                    if vec is not None:
                        r.cohort_percentiles = calibrator.compute_student_percentiles(
                            vec, distributions
                        )

                n_annotated = sum(1 for r in results if r.cohort_percentiles)
                print(f"  Cohort calibration: {len(signal_vectors)} vectors, "
                      f"{n_annotated} students annotated with class-relative percentiles")

        except Exception as _ce:
            print(f"  ⚠ Cohort calibration skipped: {_ce}")

    # Report any errors that occurred
    if errors:
        print(f"\n⚠ Warning: {len(errors)} submission(s) had errors during analysis:")
        for error in errors[:5]:  # Show first 5
            print(f"  • {error}")
        if len(errors) > 5:
            print(f"  ... and {len(errors) - 5} more")

    if not generate_report:
        return results, None

    # Generate report
    print("Generating report...")

    # Get assignment name
    assignments = get_assignments(course_id)
    assignment_name = "Unknown Assignment"
    for a in assignments:
        if a.get("id") == assignment_id:
            assignment_name = a.get("name", "Unknown Assignment")
            break

    report_gen = ReportGenerator()
    profile_name = ASSIGNMENT_PROFILES.get(profile_id, {}).get("name", profile_id)
    report_path = report_gen.generate_report(
        assignment_name=assignment_name,
        profile_name=profile_name,
        results=results,
        cohort_stats=cohort_stats,
        context_profile=context_profile
    )

    print(f"\nReport saved to: {report_path}")

    return results, report_path


# =============================================================================
# INTERACTIVE MENU
# =============================================================================

def select_from_list(items: List[Dict], prompt: str, name_key: str = "name", id_key: str = "id") -> Optional[Dict]:
    """Display a selection menu and return the chosen item."""
    if not items:
        print("No items available.")
        return None
    
    print(f"\n{prompt}")
    print("-" * 40)
    
    for i, item in enumerate(items, 1):
        name = item.get(name_key, f"Item {i}")
        print(f"  {i}. {name}")
    
    print(f"  0. Cancel")
    
    while True:
        try:
            choice = input("\nEnter number: ").strip()
            if choice == "0":
                return None
            idx = int(choice) - 1
            if 0 <= idx < len(items):
                return items[idx]
            print("Invalid selection.")
        except ValueError:
            print("Please enter a number.")


def select_profile() -> str:
    """Select an assignment profile."""
    print("\nSelect Assignment Profile:")
    print("-" * 40)
    
    profiles = list(ASSIGNMENT_PROFILES.values())
    for i, profile in enumerate(profiles, 1):
        print(f"  {i}. {profile['name']}")
        print(f"      {profile['description']}")
    
    while True:
        try:
            choice = input("\nEnter number (default 1): ").strip()
            if not choice:
                return profiles[0]['id']
            idx = int(choice) - 1
            if 0 <= idx < len(profiles):
                return profiles[idx]['id']
            print("Invalid selection.")
        except ValueError:
            print("Please enter a number.")


def main_menu():
    """Display main interactive menu."""
    print("\n" + "=" * 60)
    print("  ACADEMIC DISHONESTY CHECK v2.0")
    print("  Adaptive Academic Integrity Analysis")
    print("=" * 60)
    
    if not API_TOKEN:
        print("\n⚠️  CANVAS_API_TOKEN not set!")
        print("   Set this environment variable to connect to Canvas.")
        print("   Example: export CANVAS_API_TOKEN='your_token_here'")
        return
    
    while True:
        print("\n" + "-" * 40)
        print("MAIN MENU")
        print("-" * 40)
        print("  1. Analyze Canvas Assignment")
        print("  2. Analyze Single Text (paste)")
        print("  3. View Assignment Profiles")
        print("  4. View Feedback Statistics")
        print("  5. About Academic Integrity Detection")
        print("  6. Settings & Configuration")
        print("  0. Exit")
        
        choice = input("\nEnter choice: ").strip()
        
        if choice == "0":
            print("\nGoodbye!")
            break
        
        elif choice == "1":
            # Analyze Canvas assignment
            courses = get_courses()
            course = select_from_list(courses, "Select a Course:")
            if not course:
                continue
            
            assignments = get_assignments(course['id'])
            assignment = select_from_list(assignments, "Select an Assignment:")
            if not assignment:
                continue
            
            profile_id = select_profile()
            
            print("\nContext Profile:")
            print("  1. Community College (more lenient, ESL-aware)")
            print("  2. Standard (default thresholds)")
            ctx_choice = input("Enter choice (default 1): ").strip()
            context_profile = "community_college" if ctx_choice != "2" else "standard"
            
            results, report_path = analyze_assignment(
                course['id'],
                assignment['id'],
                profile_id=profile_id,
                context_profile=context_profile
            )
            
            if results:
                # Show summary
                high = sum(1 for r in results if r.concern_level == 'high')
                elevated = sum(1 for r in results if r.concern_level == 'elevated')
                print(f"\n📊 Summary: {len(results)} analyzed, {high} high concern, {elevated} elevated")
                
                if report_path:
                    print(f"📄 Full report: {report_path}")
        
        elif choice == "2":
            # Analyze single text
            print("\nPaste the text to analyze (enter a blank line when done):")
            lines = []
            while True:
                line = input()
                if line == "":
                    break
                lines.append(line)
            
            text = "\n".join(lines)
            if not text.strip():
                print("No text provided.")
                continue
            
            profile_id = select_profile()
            
            analyzer = DishonestyAnalyzer(profile_id=profile_id)
            result = analyzer.analyze_text(text, "manual", "Manual Input")
            
            print("\n" + "=" * 50)
            print("ANALYSIS RESULT")
            print("=" * 50)
            print(f"Concern Level: {result.concern_level.upper()}")
            print(f"Suspicious Score: {result.suspicious_score:.2f}")
            print(f"Authenticity Score: {result.authenticity_score:.2f}")
            print(f"Word Count: {result.word_count}")
            
            if result.markers_found:
                print("\nMarkers Found:")
                for marker_type, instances in result.markers_found.items():
                    print(f"  {marker_type}: {instances[:3]}")
            
            if result.conversation_starters:
                print("\nConversation Starters:")
                for starter in result.conversation_starters[:3]:
                    print(f"  • {starter}")
        
        elif choice == "3":
            # View profiles
            print("\n" + "=" * 50)
            print("ASSIGNMENT PROFILES")
            print("=" * 50)
            for profile in ASSIGNMENT_PROFILES.values():
                print(f"\n{profile['name']}")
                print(f"  {profile['description']}")
                print("  Instructor Notes:")
                for note in profile.get('instructor_notes', []):
                    print(f"    • {note}")
        
        elif choice == "4":
            # Feedback statistics
            if HAS_FEEDBACK_TRACKER:
                print_feedback_summary()
                print("\nTip: After conversations with flagged students, you can record")
                print("     outcomes to track false positive rates and marker accuracy.")
            else:
                print("\n⚠️ Feedback tracking module not available.")

        elif choice == "5":
            # About Academic Integrity Detection - Full Philosophical Framework
            show_academic_integrity_philosophy()

        elif choice == "6":
            # Settings & Configuration submenu
            show_settings_menu()
        
        else:
            print("Invalid choice.")


def show_academic_integrity_philosophy():
    """Display the philosophical framework for academic dishonesty detection."""
    print("\n" + "=" * 70)
    print("WHAT THIS TOOL DETECTS: Academic Dishonesty Defined")
    print("=" * 70)
    print("""
This tool identifies markers of ACADEMIC DISHONESTY, which means:

Using AI or other tools in ways that:
  • Compromise student authorship (AI generates content claimed as own thinking)
  • Substitute for intellectual engagement (AI replaces genuine engagement)
  • Misrepresent the origin of work (not distinguishing own ideas from AI)
""")
    input("Press Enter to continue...")
    
    print("\n" + "-" * 70)
    print("WHAT THIS TOOL IS:")
    print("-" * 70)
    print("""
  ✓ A CONVERSATION STARTER (not a verdict)
    Flags indicate submissions that warrant discussion with the student.
    They are NOT proof of dishonesty.
    
  ✓ AN OUTLIER DETECTOR (peer comparison, not absolute judgment)
    Uses statistical comparison within each class cohort.
    What's "normal" varies by class, assignment, and population.
    
  ✓ A PEDAGOGICAL SUPPORT (upholds learning objectives)
    Helps instructors identify students who may not be meeting
    learning objectives around authorship and intellectual engagement.
    
  ✓ CONTEXT-AWARE (adjusts for assignment type and student population)
    Reduces false positives for ESL students, first-generation students,
    and other diverse learners who may use patterns that "look like AI"
    but are actually legitimate learned structures.
""")
    input("Press Enter to continue...")
    
    print("\n" + "-" * 70)
    print("WHAT THIS TOOL IS NOT:")
    print("-" * 70)
    print("""
  ✗ A PROOF OF CHEATING (requires instructor judgment)
    Flags require human review and conversation with the student.
    Many flags have innocent explanations.
    
  ✗ A PUNISHMENT MECHANISM (supports teaching conversations)
    Designed to open dialogue, not to automatically penalize.
    Frame concerns as teaching opportunities.
    
  ✗ AN "AI DETECTOR" (detects dishonest use, not all use)
    Legitimate AI use exists (brainstorming, editing, etc.).
    This tool detects PATTERNS suggesting dishonesty, not AI use per se.
    
  ✗ INFALLIBLE (false positives occur; human judgment essential)
    No automated tool can perfectly distinguish AI from human text.
    Always give students the benefit of the doubt.
""")
    input("Press Enter to continue...")
    
    print("\n" + "-" * 70)
    print("PEDAGOGICAL VALUES:")
    print("-" * 70)
    print("""
This tool reflects specific commitments about teaching and learning:

  1. AUTHORSHIP IS CENTRAL
     Students must be the primary intellectual source of submitted work.
     Detection implication: Absence of authentic voice is concerning.
     
  2. LEARNING OVER PRODUCTS
     The process of intellectual engagement matters more than polish.
     Detection implication: Too-perfect work without process deserves scrutiny.
     
  3. INTELLECTUAL ENGAGEMENT REQUIRED
     Students cannot substitute AI for reading, thinking, or analyzing.
     Detection implication: Generic content suggests lack of engagement.
     
  4. CONTEXT DETERMINES DISHONESTY
     The same AI use can be legitimate or dishonest depending on assignment.
     Detection implication: Different thresholds for different assignment types.
     
  5. EQUITY AND INCLUSION
     Must account for legitimate variation in diverse student populations.
     Detection implication: Reduced weights for ESL, first-gen, and other markers.

For complete documentation, see USER_GUIDE.md in the docs/ folder.
""")
    input("Press Enter to return to menu...")


def show_settings_menu():
    """Display settings and configuration menu."""
    while True:
        print("\n" + "=" * 60)
        print("SETTINGS & CONFIGURATION")
        print("=" * 60)
        print("  1. Configure Institution Demographics")
        print("  2. Data Collection & Privacy")
        print("  3. View Stored Telemetry Data")
        print("  4. Check for Marker Updates")
        print("  0. Back to Main Menu")
        
        choice = input("\nEnter choice: ").strip()
        
        if choice == "0":
            break
        
        elif choice == "1":
            # Demographics configuration
            try:
                from modules.demographic_collector import get_demographic_collector
                collector = get_demographic_collector()
                collector.gather_demographics_interactive()
            except ImportError:
                print("\n⚠️  Demographic collector module not available.")
                print("   Demographics will use national averages.")
        
        elif choice == "2":
            # Telemetry consent
            try:
                from modules.telemetry_manager import TelemetryManager, TelemetrySystem
                manager = TelemetryManager()
                consent = manager.get_consent()
                
                print("\n" + "-" * 50)
                print("CURRENT DATA COLLECTION SETTINGS")
                print("-" * 50)
                print(f"  Program usage data:  {'ENABLED' if consent.program_usage else 'DISABLED'}")
                print(f"  Marker feedback:     {'ENABLED' if consent.marker_data else 'DISABLED'}")
                print(f"  Upload frequency:    {consent.upload_frequency.upper()}")
                
                print("\nOptions:")
                print("  1. Toggle program usage data")
                print("  2. Toggle marker feedback")
                print("  3. Change upload frequency")
                print("  0. Back")
                
                sub = input("\nEnter choice: ").strip()
                
                if sub == "1":
                    manager.update_consent(TelemetrySystem.PROGRAM_USAGE, not consent.program_usage)
                    print(f"  → Program usage data: {'ENABLED' if not consent.program_usage else 'DISABLED'}")
                elif sub == "2":
                    manager.update_consent(TelemetrySystem.MARKER_DATA, not consent.marker_data)
                    print(f"  → Marker feedback: {'ENABLED' if not consent.marker_data else 'DISABLED'}")
                elif sub == "3":
                    print("\n  1. Manual (you choose when)")
                    print("  2. Weekly")
                    print("  3. Monthly")
                    freq = input("Choose frequency: ").strip()
                    if freq == "2":
                        manager.update_consent(TelemetrySystem.PROGRAM_USAGE, consent.program_usage, "weekly")
                    elif freq == "3":
                        manager.update_consent(TelemetrySystem.PROGRAM_USAGE, consent.program_usage, "monthly")
                    else:
                        manager.update_consent(TelemetrySystem.PROGRAM_USAGE, consent.program_usage, "manual")
                    print("  → Frequency updated")
                    
            except ImportError:
                print("\n⚠️  Telemetry manager module not available.")
        
        elif choice == "3":
            # View stored data
            try:
                from modules.telemetry_manager import TelemetryManager
                manager = TelemetryManager()
                summary = manager.get_stored_data_summary()
                
                print("\n" + "-" * 50)
                print("STORED TELEMETRY DATA")
                print("-" * 50)
                print(f"  Usage events:    {summary['usage_events']['total']} ({summary['usage_events']['pending']} pending)")
                print(f"  Marker feedback: {summary['marker_feedback']['total']} ({summary['marker_feedback']['pending']} pending)")
                print(f"  Last upload:     {summary['last_upload'] or 'Never'}")
                
                if summary['usage_events']['pending'] > 0 or summary['marker_feedback']['pending'] > 0:
                    print("\nOptions:")
                    print("  1. View pending data")
                    print("  2. Delete all data")
                    print("  0. Back")
                    
                    sub = input("\nEnter choice: ").strip()
                    if sub == "2":
                        confirm = input("Are you sure? (yes/no): ").strip().lower()
                        if confirm == "yes":
                            manager.delete_all_data()
                            print("  → All telemetry data deleted")
                            
            except ImportError:
                print("\n⚠️  Telemetry manager module not available.")
        
        elif choice == "4":
            # Check for updates
            try:
                from modules.update_checker import UpdateChecker
                checker = UpdateChecker()
                result = checker.check_for_updates()
                
                print("\n" + "-" * 50)
                print("MARKER UPDATE CHECK")
                print("-" * 50)
                
                if result.updates_available:
                    print("Updates available:")
                    for update in result.available_updates:
                        print(f"  • {update['marker_id']}: {update['current_version']} → {update['new_version']}")
                    print("\nNote: Updates must be manually reviewed before applying.")
                else:
                    print("All markers are up to date.")
                    
            except ImportError:
                print("\n⚠️  Update checker module not available.")


# =============================================================================
# MAIN ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    # Check if running interactively or with arguments
    if len(sys.argv) > 1:
        # Command line mode
        if sys.argv[1] == "--help":
            print(f"""
Academic Dishonesty Check v{VERSION}

Usage:
  python {sys.argv[0]}              - Interactive menu
  python {sys.argv[0]} --help       - Show this help
  python {sys.argv[0]} --version    - Show version

Environment Variables:
  CANVAS_API_TOKEN    - Your Canvas API token (required)
  CANVAS_BASE_URL     - Canvas instance URL (e.g., https://institution.instructure.com)
""")
        elif sys.argv[1] == "--version":
            print(f"Academic Dishonesty Check v{VERSION} ({VERSION_DATE})")
        else:
            print(f"Unknown argument: {sys.argv[1]}")
            print(f"Use --help for usage information.")
    else:
        # Interactive mode
        main_menu()
