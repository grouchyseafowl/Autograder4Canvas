"""
Preprocessing pipeline orchestrator.

The insights pipeline calls this module with a list of student submissions
(already fetched from Canvas). The pipeline:

  1. Extracts text from each submission (text body + attachments)
  2. Detects audio attachments → transcribes (and translates if non-English)
  3. Detects non-English text → translates via LLM
  4. Returns PreprocessedSubmission objects with uniform English text
     ready for the insights engine

Audio translation uses Whisper's built-in --translate mode (one step,
high quality). Text translation uses chunked LLM calls. These are
separate paths because Whisper handles audio→English far better than
transcribe→LLM-translate.

The pipeline does NOT fetch submissions from Canvas — that's the
insights system's job. This module only transforms what it receives.

Usage (from insights pipeline):
    from preprocessing import PreprocessingPipeline, PreprocessedSubmission

    pipeline = PreprocessingPipeline(
        canvas_headers={"Authorization": f"Bearer {token}"},
    )
    results = pipeline.process_submissions(submissions)

    for result in results:
        if result.text:
            # Feed to insights engine
            ...
        if result.needs_teacher_comment:
            # Post translation/transcription as Canvas comment
            ...
"""

import logging
import os
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

import requests as http_requests

from .image_transcriber import ImageTranscriber, ImageTranscriptionResult, is_image_attachment
from .language_detector import detect_language, language_name, LanguageResult
from .transcriber import Transcriber, TranscriptionResult, is_audio_attachment
from .translator import Translator, TranslationResult

logger = logging.getLogger("autograder.preprocessing")


# --- Text extraction helpers (mirrors existing autograder logic) ---

def _clean_html(text: str) -> str:
    """Strip HTML tags."""
    return re.sub(r'<[^>]+>', ' ', text or "").strip()


def _extract_text_attachment(attachment: Dict, headers: Dict) -> Optional[str]:
    """
    Extract text from a non-audio/image attachment.

    Supported: .txt, .md, .html, .htm, .docx, .doc, .odt, .pdf, .rtf

    Replicates the extraction logic from the main autograder so the
    preprocessing pipeline can work independently.
    """
    import requests
    filename = attachment.get("filename", "")
    ext = Path(filename).suffix.lower()
    url = attachment.get("url", "")

    if not url:
        return None

    try:
        r = requests.get(url, headers=headers, timeout=60)
        r.raise_for_status()
        content = r.content

        # Plain text / Markdown — treat as-is
        if ext in (".txt", ".md"):
            return content.decode("utf-8", errors="replace")

        # HTML — strip tags
        elif ext in (".html", ".htm"):
            import re
            raw = content.decode("utf-8", errors="replace")
            return re.sub(r"<[^>]+>", " ", raw).strip()

        elif ext == ".docx":
            try:
                import io
                from docx import Document
                doc = Document(io.BytesIO(content))
                return "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                logger.warning("python-docx not available for .docx extraction")
                return None

        elif ext == ".doc":
            # Old binary Word format — python-docx sometimes handles simple .doc files
            try:
                import io
                from docx import Document
                doc = Document(io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if text:
                    return text
            except Exception:
                pass
            logger.warning(
                "Could not extract text from .doc file '%s' — "
                "old binary Word format. Convert to .docx for reliable extraction.",
                filename,
            )
            return None

        elif ext == ".odt":
            try:
                import io
                from odf.opendocument import load as odf_load
                from odf.text import P
                from odf import teletype
                doc = odf_load(io.BytesIO(content))
                paragraphs = doc.getElementsByType(P)
                lines = []
                for p in paragraphs:
                    line = teletype.extractText(p)
                    if line.strip():
                        lines.append(line)
                return "\n".join(lines)
            except ImportError:
                logger.warning(
                    "odfpy not installed — cannot extract text from ODT '%s'. "
                    "Install with: pip install odfpy",
                    filename,
                )
                return None

        elif ext == ".pdf":
            try:
                import io
                from pdfminer.high_level import extract_text
                return extract_text(io.BytesIO(content))
            except ImportError:
                logger.warning(
                    "pdfminer.six not installed — cannot extract text from PDF '%s'. "
                    "Install with: pip install pdfminer.six",
                    filename,
                )
                return None

        elif ext == ".rtf":
            try:
                from striprtf.striprtf import rtf_to_text
                return rtf_to_text(content.decode("utf-8", errors="replace"))
            except ImportError:
                logger.warning(
                    "striprtf not installed — cannot extract text from RTF '%s'. "
                    "Install with: pip install striprtf",
                    filename,
                )
                return None

        return None

    except Exception as e:
        logger.warning(f"Failed to extract text from {filename}: {e}")
        return None


# --- Result types ---

@dataclass
class PreprocessedSubmission:
    """
    A student submission after preprocessing.

    The insights pipeline consumes `text` — which is always English
    (either original English, or translated/transcribed).
    """
    # Identifiers (passed through from Canvas submission)
    submission_id: int
    user_id: int
    assignment_id: int

    # The final English text for the insights engine
    text: str

    # What processing was done
    was_translated: bool = False
    was_transcribed: bool = False
    was_image_transcribed: bool = False
    original_language: Optional[str] = None
    original_language_name: Optional[str] = None
    # Three-way multilingual classification from language_detector
    multilingual_type: str = "monolingual_english"
    detected_languages: Optional[List[str]] = None

    # Original content preserved for reference
    original_text: Optional[str] = None
    transcription_results: List[TranscriptionResult] = field(default_factory=list)
    image_transcription_results: List[ImageTranscriptionResult] = field(default_factory=list)
    translation_result: Optional[TranslationResult] = None

    # Teacher comment content (if processing was done, teacher may want to know)
    teacher_comment: Optional[str] = None

    @property
    def needs_teacher_comment(self) -> bool:
        """True if preprocessing produced content worth noting for the teacher."""
        return self.teacher_comment is not None

    @property
    def processing_summary(self) -> str:
        """One-line summary of what was done."""
        parts = []
        if self.was_transcribed:
            count = len(self.transcription_results)
            translated_audio = sum(
                1 for r in self.transcription_results if r.was_translated
            )
            desc = f"transcribed {count} audio file(s)"
            if translated_audio:
                desc += f" ({translated_audio} translated to English)"
            parts.append(desc)
        if self.was_image_transcribed:
            count = len(self.image_transcription_results)
            parts.append(
                f"transcribed {count} handwritten image(s) "
                f"[NEEDS VERIFICATION]"
            )
        if self.was_translated:
            parts.append(f"text translated from {self.original_language_name}")
        if not parts:
            return "no preprocessing needed"
        return ", ".join(parts)


class PreprocessingPipeline:
    """
    Orchestrates language detection, translation, and audio transcription
    for a batch of student submissions.
    """

    def __init__(
        self,
        canvas_headers: Optional[Dict[str, str]] = None,
        # Translation config
        translation_enabled: bool = True,
        translation_backend: str = "ollama",
        translation_model: str = "llama3.1:8b",
        ollama_base_url: str = "http://localhost:11434",
        cloud_api_key: Optional[str] = None,
        cloud_base_url: Optional[str] = None,
        cloud_model: Optional[str] = None,
        # Transcription config
        transcription_enabled: bool = True,
        whisper_backend: str = "auto",
        whisper_model: str = "base",
        whisper_device: str = "auto",
        whisper_cpp_binary: Optional[str] = None,
        whisper_cpp_model: Optional[str] = None,
        whisper_cpp_threads: int = 4,
        # Image transcription config (handwritten notes)
        image_transcription_enabled: bool = False,
        image_transcription_backend: str = "ollama",
        image_transcription_model: str = "llama3.2-vision:11b",
        image_dpi: int = 150,  # 150 dpi is the sweet spot for small vision models
        # Behavior
        generate_teacher_comments: bool = True,
    ):
        self.canvas_headers = canvas_headers or {}
        self.translation_enabled = translation_enabled
        self.transcription_enabled = transcription_enabled
        self.image_transcription_enabled = image_transcription_enabled
        self.image_dpi = image_dpi
        self.generate_teacher_comments = generate_teacher_comments

        self._translator = None
        self._transcriber = None
        self._image_transcriber = None

        # Store config for lazy init
        self._image_transcription_config = {
            "backend": image_transcription_backend,
            "model": image_transcription_model,
            "ollama_base_url": ollama_base_url,
        }
        self._translation_config = {
            "backend": translation_backend,
            "model": translation_model,
            "base_url": ollama_base_url,
            "api_key": cloud_api_key,
            "cloud_base_url": cloud_base_url,
            "cloud_model": cloud_model,
        }
        self._transcription_config = {
            "backend": whisper_backend,
            "model_size": whisper_model,
            "device": whisper_device,
            "whisper_cpp_binary": whisper_cpp_binary,
            "whisper_cpp_model": whisper_cpp_model,
            "whisper_cpp_threads": whisper_cpp_threads,
        }

    @property
    def translator(self) -> Translator:
        if self._translator is None:
            self._translator = Translator(**self._translation_config)
        return self._translator

    @property
    def transcriber(self) -> Transcriber:
        if self._transcriber is None:
            self._transcriber = Transcriber(**self._transcription_config)
        return self._transcriber

    @property
    def image_transcriber(self) -> ImageTranscriber:
        if self._image_transcriber is None:
            self._image_transcriber = ImageTranscriber(
                **self._image_transcription_config
            )
        return self._image_transcriber

    def _extract_submission_text(self, submission: Dict) -> str:
        """Pull text from submission body + text attachments (may be any language)."""
        parts = []

        body = submission.get("body")
        if body:
            parts.append(_clean_html(body))

        for att in submission.get("attachments", []):
            if not is_audio_attachment(att) and not is_image_attachment(att):
                text = _extract_text_attachment(att, self.canvas_headers)
                if text:
                    parts.append(text.strip())

        return "\n\n".join(parts)

    def _get_audio_attachments(self, submission: Dict) -> List[Dict]:
        return [
            att for att in submission.get("attachments", [])
            if is_audio_attachment(att)
        ]

    def _transcribe_audio(
        self, audio_attachments: List[Dict], sub_id: int,
    ) -> List[TranscriptionResult]:
        """
        Transcribe audio attachments.

        Downloads each file once, then:
          1. Transcribe to detect language
          2. If non-English, re-transcribe with --translate (no re-download)

        Whisper's translate mode outputs English directly in one model
        pass — far better quality than transcribe → LLM translate.
        """
        results = []

        for att in audio_attachments:
            filename = att.get("filename", "audio")
            url = att.get("url", "")
            ext = Path(filename).suffix.lower() or ".wav"
            tmp_path = None

            try:
                # Download once
                logger.info(f"  Downloading audio: {filename}...")
                r = http_requests.get(
                    url, headers=self.canvas_headers,
                    timeout=120, stream=True,
                )
                r.raise_for_status()

                with tempfile.NamedTemporaryFile(
                    suffix=ext, delete=False
                ) as tmp:
                    for chunk in r.iter_content(chunk_size=8192):
                        tmp.write(chunk)
                    tmp_path = tmp.name

                # First pass: transcribe in original language
                logger.info(f"  Transcribing: {filename}...")
                result = self.transcriber.transcribe_file(
                    tmp_path, translate=False,
                )
                result.filename = filename

                if not result.success:
                    results.append(result)
                    continue

                # If non-English, re-transcribe with translation (same file)
                if (result.detected_language
                        and result.detected_language != "en"
                        and result.transcript.strip()):
                    logger.info(
                        f"  Audio detected as {result.detected_language} "
                        f"— re-transcribing with Whisper translate mode..."
                    )
                    translated = self.transcriber.transcribe_file(
                        tmp_path, translate=True,
                    )
                    translated.filename = filename
                    if translated.success:
                        translated.detected_language = result.detected_language
                        results.append(translated)
                        continue

                    logger.warning(
                        f"  Whisper translate failed — keeping "
                        f"original-language transcription"
                    )

                results.append(result)

            except Exception as e:
                logger.error(
                    f"  Download failed for {filename}: {e}"
                )
                results.append(TranscriptionResult(
                    filename=filename, transcript="",
                    success=False, error=str(e),
                ))
            finally:
                if tmp_path:
                    try:
                        os.unlink(tmp_path)
                    except OSError:
                        pass

        return results

    def _build_teacher_comment(
        self,
        transcription_results: List[TranscriptionResult],
        image_transcription_results: Optional[List[ImageTranscriptionResult]] = None,
        translation_result: Optional[TranslationResult] = None,
    ) -> Optional[str]:
        """Build an informational teacher comment noting what preprocessing was done."""
        if not self.generate_teacher_comments:
            return None

        lines = ["[Autograder — Preprocessing Note]", ""]

        if transcription_results:
            lines.append("Audio Transcription:")
            for tr in transcription_results:
                if tr.success:
                    mins = int(tr.duration_seconds // 60)
                    secs = int(tr.duration_seconds % 60)
                    parts = [f"{tr.filename}: {mins}:{secs:02d}"]
                    if tr.detected_language and tr.detected_language != "en":
                        parts.append(
                            f"detected {tr.detected_language}"
                            + (" → translated to English" if tr.was_translated else "")
                        )
                    lines.append(f"  - {', '.join(parts)}")
                else:
                    lines.append(f"  - {tr.filename}: failed ({tr.error})")
            lines.append("")

        if image_transcription_results:
            has_text = any(r.success and r.transcript for r in image_transcription_results)
            has_visual = any(
                r.success and getattr(r, "description", "") for r in image_transcription_results
            )

            if has_text:
                lines.append("Image Submission — Text Content (NEEDS VERIFICATION):")
                lines.append(
                    "  ⚠ Text was extracted from the submitted image(s) using AI. "
                    "Please verify accuracy before relying on this content."
                )
                for ir in image_transcription_results:
                    if ir.success and ir.transcript:
                        preview = ir.transcript[:100]
                        if len(ir.transcript) > 100:
                            preview += "..."
                        lines.append(f"  - {ir.filename}: \"{preview}\"")
                    elif not ir.success:
                        lines.append(f"  - {ir.filename}: failed ({ir.error})")
                lines.append("")

            if has_visual:
                lines.append("Image Submission — Visual Description (AI-generated):")
                lines.append(
                    "  ⚠ These are AI descriptions of visual/artistic content. "
                    "View the original file to assess the submission fully."
                )
                for ir in image_transcription_results:
                    desc = getattr(ir, "description", "")
                    if ir.success and desc and not ir.transcript:
                        lines.append(f"  - {ir.filename}: {desc}")
                lines.append("")

        if translation_result and translation_result.success:
            lines.append(
                f"Text translated from "
                f"{translation_result.source_language_name} to English "
                f"({translation_result.chunks_translated} chunk(s))"
            )
            lines.append("")

        if len(lines) <= 2:
            return None

        return "\n".join(lines)

    def process_submission(self, submission: Dict) -> PreprocessedSubmission:
        """
        Process a single Canvas submission.

        Processing flow:
          1. Transcribe audio → get English text (Whisper translates if needed)
          2. Extract text content from body/attachments
          3. Detect language on TEXT only (audio already handled by Whisper)
          4. Translate text via LLM if non-English
          5. Combine all English text
        """
        sub_id = submission.get("id", 0)
        user_id = submission.get("user_id", 0)
        assignment_id = submission.get("assignment_id", 0)

        # --- Step 1: Audio transcription ---
        transcription_results = []
        transcribed_texts = []

        if self.transcription_enabled:
            audio_attachments = self._get_audio_attachments(submission)
            if audio_attachments:
                if not self.transcriber.is_available():
                    logger.warning(
                        f"  Submission {sub_id}: {len(audio_attachments)} "
                        f"audio file(s) but no transcription backend available"
                    )
                else:
                    transcription_results = self._transcribe_audio(
                        audio_attachments, sub_id
                    )
                    for tr in transcription_results:
                        if tr.success and tr.transcript:
                            transcribed_texts.append(tr.transcript)

        # --- Step 1b: Image transcription (handwritten notes) ---
        image_transcription_results = []
        image_texts = []

        if self.image_transcription_enabled:
            image_attachments = [
                att for att in submission.get("attachments", [])
                if is_image_attachment(att)
            ]
            if image_attachments:
                if not self.image_transcriber.is_available():
                    logger.warning(
                        f"  Submission {sub_id}: {len(image_attachments)} "
                        f"image file(s) but no vision model available"
                    )
                else:
                    for att in image_attachments:
                        fname = att.get("filename", "image")
                        url = att.get("url", "")
                        if url:
                            logger.info(f"  Transcribing handwritten: {fname}...")
                            result = self.image_transcriber.transcribe_from_url(
                                url, fname, headers=self.canvas_headers,
                            )
                            image_transcription_results.append(result)
                            if result.success and result.transcript:
                                image_texts.append(result.transcript)

        # --- Step 2: Extract text content ---
        text_content = self._extract_submission_text(submission)

        # --- Step 3: Language detection on TEXT only ---
        # Audio is already handled by Whisper's translate mode.
        # We only need LLM translation for the text portion.
        translation_result = None
        translated_text = text_content
        lang_result = None

        if self.translation_enabled and text_content.strip():
            lang_result = detect_language(text_content)

            if lang_result.needs_translation:
                logger.info(
                    f"  Submission {sub_id}: text detected as "
                    f"{language_name(lang_result.language_code)} "
                    f"({lang_result.confidence:.0%}) — translating..."
                )
                translation_result = self.translator.translate(
                    text_content, lang_result.language_code
                )
                if translation_result.success:
                    translated_text = translation_result.translated_text

        # --- Step 4: Combine all English text ---
        final_parts = []
        if translated_text.strip():
            final_parts.append(translated_text)
        final_parts.extend(transcribed_texts)
        # Image transcriptions — add clean text (no annotations that would
        # confuse the LLM). The was_image_transcribed flag tells the UI to
        # display verification warnings.
        final_parts.extend(image_texts)
        final_text = "\n\n".join(final_parts)

        if not final_text.strip():
            return PreprocessedSubmission(
                submission_id=sub_id,
                user_id=user_id,
                assignment_id=assignment_id,
                text="",
            )

        # --- Step 5: Build teacher comment ---
        teacher_comment = self._build_teacher_comment(
            transcription_results, image_transcription_results,
            translation_result,
        )

        was_transcribed = any(r.success for r in transcription_results)
        was_image_transcribed = any(
            r.success and r.transcript for r in image_transcription_results
        )
        was_translated = (
            translation_result is not None and translation_result.success
        )

        # Preserve original text if anything was transformed
        original_text = None
        if was_translated or was_transcribed:
            original_parts = []
            if text_content.strip():
                original_parts.append(text_content)
            original_parts.extend(
                r.transcript for r in transcription_results
                if r.success and not r.was_translated
            )
            original_text = "\n\n".join(original_parts) if original_parts else None

        return PreprocessedSubmission(
            submission_id=sub_id,
            user_id=user_id,
            assignment_id=assignment_id,
            text=final_text,
            was_translated=was_translated,
            was_transcribed=was_transcribed,
            was_image_transcribed=was_image_transcribed,
            original_language=(
                translation_result.source_language if translation_result else None
            ),
            original_language_name=(
                translation_result.source_language_name if translation_result else None
            ),
            original_text=original_text,
            transcription_results=transcription_results,
            image_transcription_results=image_transcription_results,
            translation_result=translation_result,
            teacher_comment=teacher_comment,
            multilingual_type=(
                lang_result.multilingual_type if lang_result else "monolingual_english"
            ),
            detected_languages=(
                lang_result.detected_languages if lang_result else None
            ),
        )

    def process_submissions(
        self,
        submissions: List[Dict],
        progress_callback=None,
    ) -> List[PreprocessedSubmission]:
        """
        Process a batch of Canvas submissions.

        Args:
            submissions: List of Canvas submission dicts.
            progress_callback: Optional callable(current, total, message)
                for progress reporting (e.g. to GUI).

        Returns:
            List of PreprocessedSubmission, one per input submission.
        """
        results = []
        total = len(submissions)

        for i, sub in enumerate(submissions):
            user_id = sub.get("user_id", "?")

            if progress_callback:
                progress_callback(
                    i, total,
                    f"Preprocessing submission {i+1}/{total} "
                    f"(student {user_id})"
                )

            logger.info(f"Processing submission {sub.get('id', '?')} ({i+1}/{total})...")
            result = self.process_submission(sub)

            if result.was_transcribed or result.was_translated:
                logger.info(f"  → {result.processing_summary}")

            results.append(result)

        if progress_callback:
            progress_callback(total, total, "Preprocessing complete")

        translated_count = sum(1 for r in results if r.was_translated)
        transcribed_count = sum(1 for r in results if r.was_transcribed)
        if translated_count or transcribed_count:
            logger.info(
                f"Preprocessing complete: {translated_count} translated, "
                f"{transcribed_count} transcribed out of {total} submissions"
            )

        return results
